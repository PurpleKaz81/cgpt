#!/usr/bin/env python3
import argparse
import json
import os
import re
import shutil
import sys
import zipfile
import sqlite3
import time
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

try:
    from zoneinfo import ZoneInfo  # py3.9+
except Exception:
    ZoneInfo = None  # fallback


# Enable line-editing for interactive `input()` (arrow keys, history, tab completion).
# On macOS this typically wraps libedit; ignore failures if module/bindings differ.
try:
    import readline  # noqa: F401

    try:
        # Bind tab completion if available; some libedit builds may raise here.
        readline.parse_and_bind("tab: complete")
    except Exception:
        pass
except Exception:
    pass

__version__ = "0.2.0"

SAO_PAULO_TZ = "America/Sao_Paulo"


# ----------------- util -----------------


def die(msg: str, code: int = 1) -> None:
    print(f"ERROR: {msg}", file=sys.stderr)
    sys.exit(code)


def looks_like_home(p: Path) -> bool:
    return (
        (p / "zips").is_dir()
        and (p / "extracted").is_dir()
        and (p / "dossiers").is_dir()
    )


def discover_home() -> Path:
    """
    If user runs cgpt.py from a random folder, try to find the chatgpt_chat_exports
    home by walking up from:
        - current working directory
        - the directory containing cgpt.py
    """
    candidates = []
    try:
        candidates.append(Path.cwd().resolve())
    except Exception:
        pass
    try:
        candidates.append(Path(__file__).resolve().parent)
    except Exception:
        pass

    seen = set()
    for base in candidates:
        for p in [base, *list(base.parents)[:8]]:
            if p in seen:
                continue
            seen.add(p)
            if looks_like_home(p):
                return p

    # Fallback: cwd (will error later if layout missing)
    return Path.cwd().resolve()


def home_dir(cli_home: Optional[str]) -> Path:
    if cli_home:
        return Path(cli_home).expanduser().resolve()
    env = os.environ.get("CGPT_HOME")
    if env:
        return Path(env).expanduser().resolve()
    return discover_home()


def ensure_layout(home: Path) -> Tuple[Path, Path, Path]:
    zips_dir = home / "zips"
    extracted_dir = home / "extracted"
    dossiers_dir = home / "dossiers"
    for d in (zips_dir, extracted_dir, dossiers_dir):
        if not d.exists():
            die(
                f"Missing folder: {d}\nExpected: zips/, extracted/, dossiers/ under {home}"
            )
        if not d.is_dir():
            die(f"Not a directory: {d}")
    return zips_dir, extracted_dir, dossiers_dir


def init_layout(home: Path) -> Tuple[List[Path], List[Path]]:
    """Create required folders under home when missing; verify existing layout."""
    required = [home / "zips", home / "extracted", home / "dossiers"]
    created: List[Path] = []
    existing: List[Path] = []

    if home.exists() and not home.is_dir():
        die(f"Home path is not a directory: {home}")
    home.mkdir(parents=True, exist_ok=True)

    for d in required:
        if d.exists():
            if not d.is_dir():
                die(f"Expected directory but found file: {d}")
            existing.append(d)
            continue
        d.mkdir(parents=True, exist_ok=True)
        created.append(d)

    return created, existing


def newest_zip(zips_dir: Path) -> Path:
    zips = sorted(zips_dir.glob("*.zip"), key=lambda p: p.stat().st_mtime, reverse=True)
    if not zips:
        die(f"No ZIPs found in {zips_dir}")
    return zips[0]


def newest_extracted(extracted_dir: Path) -> Path:
    dirs = [p for p in extracted_dir.iterdir() if p.is_dir() and p.name != "latest"]
    if not dirs:
        die(f"No extracted folders found in {extracted_dir}. Run: cgpt.py extract")
    dirs.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return dirs[0]


def refresh_latest_symlink(extracted_dir: Path, target: Path) -> None:
    """Point extracted/latest to target. Fall back to extracted/LATEST.txt if symlink fails."""
    latest = extracted_dir / "latest"
    try:
        if latest.is_symlink() or latest.is_file():
            latest.unlink()
        elif latest.is_dir():
            shutil.rmtree(latest)
        latest.symlink_to(target, target_is_directory=True)
        # also keep pointer file as a backup
        (extracted_dir / "LATEST.txt").write_text(str(target) + "\n", encoding="utf-8")
    except OSError:
        (extracted_dir / "LATEST.txt").write_text(str(target) + "\n", encoding="utf-8")


def default_root(extracted_dir: Path) -> Path:
    """
    Prefer extracted/latest if present. Otherwise use extracted/LATEST.txt. Otherwise newest extracted dir.
    """
    latest = extracted_dir / "latest"
    if latest.exists():
        try:
            return latest.resolve()
        except Exception:
            return latest

    ptr = extracted_dir / "LATEST.txt"
    if ptr.exists():
        t = ptr.read_text(encoding="utf-8").strip()
        if t:
            p = Path(t).expanduser()
            if p.exists():
                return p.resolve()

    return newest_extracted(extracted_dir)


def safe_slug(s: str, max_len: int = 80) -> str:
    s = (s or "").strip()
    s = re.sub(r"\s+", " ", s)
    s = re.sub(r"[^\w\-\.\s]", "", s, flags=re.UNICODE)
    s = s.strip().replace(" ", "_")
    return s[:max_len] if len(s) > max_len else s


def ts_to_local_str(ts: float) -> str:
    if not ts:
        return ""
    dt_utc = datetime.fromtimestamp(ts, tz=timezone.utc)
    if ZoneInfo:
        try:
            dt_loc = dt_utc.astimezone(ZoneInfo(SAO_PAULO_TZ))
            return dt_loc.isoformat()
        except Exception:
            return dt_utc.isoformat()
    return dt_utc.isoformat()


def normalize_text(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r"\s+", " ", s)
    return s


# ----------------- color helpers -----------------
_CLI_COLOR_OVERRIDE: Optional[bool] = None


def _parse_env_bool(name: str) -> Optional[bool]:
    """Parse a boolean-like env var, returning True/False/None (invalid or unset)."""
    env = os.environ.get(name)
    if env is None:
        return None
    v = env.strip().lower()
    if v in ("1", "true", "yes", "on"):
        return True
    if v in ("0", "false", "no", "off"):
        return False
    return None


def _supports_color() -> bool:
    """Return whether to use ANSI colors.

    Honor the `CGPT_FORCE_COLOR` environment variable when present:
      - true values:  '1', 'true', 'yes', 'on'  => force enable
      - false values: '0', 'false', 'no', 'off'  => force disable
    Otherwise, fall back to `sys.stdout.isatty()`.
    """
    # CLI override takes precedence
    if _CLI_COLOR_OVERRIDE is not None:
        return _CLI_COLOR_OVERRIDE

    force_color = _parse_env_bool("CGPT_FORCE_COLOR")
    if force_color is not None:
        return force_color
    try:
        return sys.stdout.isatty()
    except Exception:
        return False


def _colorize_title_with_topic(title: str, topic: str) -> str:
    """Wrap the title in white and highlight case-insensitive `topic` matches in red.

    Falls back to the plain title when coloring isn't supported.
    """
    if not topic:
        return title
    if not _supports_color():
        return title
    try:
        red = "\033[31m"
        white = "\033[97m"
        reset = "\033[0m"
        pat = re.compile(re.escape(topic), re.IGNORECASE)
        highlighted = pat.sub(lambda m: f"{red}{m.group(0)}{white}", title)
        return f"{white}{highlighted}{reset}"
    except Exception:
        return title


def _colorize_title_with_topics(title: str, topics: List[str]) -> str:
    if not topics:
        return title
    if not _supports_color():
        return title
    try:
        red = "\033[31m"
        white = "\033[97m"
        reset = "\033[0m"
        parts = [re.escape(t) for t in topics if t]
        if not parts:
            return title
        pat = re.compile("(" + "|".join(parts) + ")", re.IGNORECASE)
        highlighted = pat.sub(lambda m: f"{red}{m.group(0)}{white}", title)
        return f"{white}{highlighted}{reset}"
    except Exception:
        return title


def find_conversations_json(root: Path) -> Optional[Path]:
    candidates = list(root.rglob("conversations*.json"))
    if candidates:
        candidates.sort(key=lambda p: p.stat().st_size, reverse=True)
        return candidates[0]
    json_files = list(root.rglob("*.json"))
    if not json_files:
        return None
    json_files.sort(key=lambda p: p.stat().st_size, reverse=True)
    return json_files[0]


def load_json(p: Path) -> Any:
    try:
        with p.open("r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        die(f"Failed to parse JSON: {p}\n{e}")


def normalize_conversations(data: Any) -> List[Dict[str, Any]]:
    if isinstance(data, list):
        return [c for c in data if isinstance(c, dict)]
    if isinstance(data, dict):
        if isinstance(data.get("conversations"), list):
            return [c for c in data["conversations"] if isinstance(c, dict)]
        if data and all(isinstance(v, dict) for v in data.values()):
            out = []
            for k, v in data.items():
                if isinstance(v, dict):
                    merged = dict(v)
                    merged.setdefault("id", k)
                    out.append(merged)
            return out
    return []


def conv_id_and_title(c: Dict[str, Any]) -> Tuple[Optional[str], str]:
    cid = c.get("id") or c.get("conversation_id") or c.get("uuid")
    title = (
        (c.get("title") or c.get("name") or "")
        .replace("\t", " ")
        .replace("\n", " ")
        .strip()
    )
    return cid, title


# ----------------- message extraction -----------------


def render_content(content: Any) -> str:
    if not isinstance(content, dict):
        return ""
    ctype = content.get("content_type")
    parts = content.get("parts")
    if ctype == "text" and isinstance(parts, list):
        return "\n".join(str(p) for p in parts if p is not None).strip()
    if isinstance(parts, list) and parts:
        return "\n".join(str(p) for p in parts if p is not None).strip()
    if "text" in content and isinstance(content["text"], str):
        return content["text"].strip()
    return ""


@dataclass
class Msg:
    t: float
    role: str
    text: str


def extract_messages_best_effort(c: Dict[str, Any]) -> List[Msg]:
    msgs: List[Msg] = []

    mapping = c.get("mapping")
    if not isinstance(mapping, dict):
        flat = c.get("messages")
        if isinstance(flat, list):
            for m in flat:
                if not isinstance(m, dict):
                    continue
                t = float(m.get("create_time") or 0.0)
                role = (m.get("author") or {}).get("role") or "unknown"
                text = render_content((m.get("content") or {})).strip()
                if text:
                    msgs.append(Msg(t=t, role=role, text=text))
        msgs.sort(key=lambda x: x.t)
        return msgs

    for node in mapping.values():
        if not isinstance(node, dict):
            continue
        m = node.get("message")
        if not isinstance(m, dict):
            continue
        t = float(m.get("create_time") or 0.0)
        author = m.get("author") or {}
        role = author.get("role") or "unknown"
        content = m.get("content") or {}
        text = render_content(content).strip()
        if text:
            msgs.append(Msg(t=t, role=role, text=text))

    msgs.sort(key=lambda x: x.t)
    return msgs


def conversation_matches_text(c: Dict[str, Any], pat: re.Pattern) -> bool:
    """
    Return True if the compiled regex `pat` matches any message text or the conversation title.
    Uses the same message extraction logic as extract_messages_best_effort.
    """
    try:
        # Check messages
        msgs = extract_messages_best_effort(c)
        for m in msgs:
            if pat.search(m.text):
                return True
        # Fallback: check title/name fields
        title = c.get("title") or c.get("name") or ""
        if isinstance(title, str) and pat.search(title):
            return True
    except Exception:
        return False
    return False


# ----------------- sqlite FTS index (simple skeleton) -----------------
def _init_index(db_path: Path) -> None:
    conn = sqlite3.connect(str(db_path))
    try:
        cur = conn.cursor()
        cur.execute(
            "CREATE TABLE IF NOT EXISTS conv_meta (id TEXT PRIMARY KEY, title TEXT, create_time REAL)"
        )
        try:
            cur.execute(
                "CREATE VIRTUAL TABLE IF NOT EXISTS conv_search USING fts5(title, content, cid UNINDEXED)"
            )
        except sqlite3.OperationalError:
            # FTS5 not available in this sqlite build
            print(
                "WARNING: SQLite FTS5 not available; full-text index disabled",
                file=sys.stderr,
            )
        conn.commit()
    finally:
        conn.close()


def index_export(
    root: Path, db_path: Path, reindex: bool = False, show_progress: bool = True
) -> None:
    """Index conversations under `root` into `db_path`.

    If `reindex` is True the existing indexed rows will be cleared first.
    This is a lightweight, best-effort implementation intended to make searches
    faster for typical-sized exports.
    """
    data_file = find_conversations_json(root)
    if not data_file:
        return
    data = load_json(data_file)
    convs = normalize_conversations(data)

    db_path.parent.mkdir(parents=True, exist_ok=True)
    _init_index(db_path)
    conn = sqlite3.connect(str(db_path))
    try:
        cur = conn.cursor()
        if reindex:
            try:
                cur.execute("DELETE FROM conv_meta")
            except Exception:
                pass
            try:
                cur.execute("DELETE FROM conv_search")
            except Exception:
                pass

        conn.execute("BEGIN")
        total = len(convs)
        i = 0
        start = time.time()
        last_update = start
        spinner = ["|", "/", "-", "\\"]
        spin_idx = 0

        def _fmt_eta(seconds: float) -> str:
            s = int(seconds + 0.5)
            m, sec = divmod(s, 60)
            h, m = divmod(m, 60)
            if h:
                return f"{h}:{m:02d}:{sec:02d}"
            return f"{m:02d}:{sec:02d}"

        for c in convs:
            cid, title = conv_id_and_title(c)
            if not cid:
                continue
            ctime = float(c.get("create_time") or 0.0)
            msgs = extract_messages_best_effort(c)
            content = "\n".join(m.text for m in msgs)
            try:
                cur.execute(
                    "REPLACE INTO conv_meta (id, title, create_time) VALUES (?, ?, ?)",
                    (cid, title, ctime),
                )
                # keep conv_search in sync: remove any prior entries for this cid
                try:
                    cur.execute("DELETE FROM conv_search WHERE cid = ?", (cid,))
                    cur.execute(
                        "INSERT INTO conv_search (title, content, cid) VALUES (?, ?, ?)",
                        (title, content, cid),
                    )
                except sqlite3.OperationalError:
                    # FTS table not available; ignore silently
                    pass
            except Exception:
                # best-effort: skip problematic conversations
                continue
            i += 1
            if show_progress:
                now = time.time()
                # update display at most every 0.25s, or on final item
                if i == total or (now - last_update >= 0.25):
                    elapsed = now - start if now > start else 0.0
                    rate = (i / elapsed) if elapsed > 0 else None
                    if rate and rate > 0:
                        rem = total - i
                        eta = rem / rate
                        eta_str = _fmt_eta(eta)
                    else:
                        eta_str = "--:--"
                    spin = spinner[spin_idx % len(spinner)]
                    spin_idx += 1
                    last_update = now
                    # carriage-return line to update progress in-place
                    print(
                        f"\rIndexed {i}/{total} {spin} ETA: {eta_str}",
                        end="",
                        file=sys.stderr,
                        flush=True,
                    )
        if show_progress:
            # finish line
            print("", file=sys.stderr)
        conn.commit()
    finally:
        conn.close()


def query_index(db_path: Path, q: str, where: str = "all") -> List[Tuple[str, str]]:
    """Query the index and return list of (cid, title).

    Falls back to empty list on any error.
    """
    if not db_path.exists():
        return []
    conn = sqlite3.connect(str(db_path))
    try:
        cur = conn.cursor()
        try:
            if where == "title":
                cur.execute(
                    "SELECT cid, title FROM conv_search WHERE title MATCH ?", (q,)
                )
            elif where == "messages":
                cur.execute(
                    "SELECT cid, title FROM conv_search WHERE content MATCH ?", (q,)
                )
            else:
                cur.execute(
                    "SELECT cid, title FROM conv_search WHERE conv_search MATCH ?", (q,)
                )
            rows = cur.fetchall()
            return [(r[0], r[1]) for r in rows]
        except sqlite3.OperationalError:
            return []
    finally:
        conn.close()


def build_fts_query(terms: List[str], and_terms: bool) -> str:
    """Build a simple FTS5 MATCH query from terms.

    Wrap each term in double-quotes and join with AND/OR.
    """
    parts: List[str] = []
    for t in terms:
        if not t:
            continue
        # escape double quotes inside term
        tp = t.replace('"', '""')
        parts.append(f'"{tp}"')
    if not parts:
        return ""
    return (" AND ".join(parts)) if and_terms else (" OR ".join(parts))


# ----------------- clean TXT formatting for AI readability -----------------


def _extract_sources(text: str) -> List[Tuple[str, str]]:
    """Extract URLs from text and return list of (url, normalized_label)."""
    url_pattern = re.compile(r"https?://[^\s\)\]\}\"\']*[^\s\)\]\}\"\'\.,;:!?]")
    matches = url_pattern.findall(text)
    unique = []
    seen = set()
    for url in matches:
        if url not in seen:
            seen.add(url)
            # normalize: remove common trailing punctuation
            url = re.sub(r"[.,;:!?\'\"]$", "", url)
            # generate label from domain + path
            try:
                from urllib.parse import urlparse

                parsed = urlparse(url)
                label = parsed.netloc + (parsed.path[:50] if parsed.path else "")
            except Exception:
                label = url[:60]
            unique.append((url, label))
    return unique


def _generate_toc(groups: Dict[str, List[Dict[str, Any]]]) -> List[str]:
    """Generate a machine-readable table of contents from conversation groups.

    Returns list of lines: title, conversation count, approximate section info.
    """
    toc: List[str] = []
    toc.append("## TABLE OF CONTENTS\n")
    line_num = 10  # rough estimate (header + metadata takes ~10 lines)

    for group_name, items in groups.items():
        root = items[0]
        root_title = root["title"] or "Untitled"
        branch_count = len(items) - 1
        create_time = root["ctime"]

        section_label = f"Line ~{line_num}: {root_title}"
        if branch_count > 0:
            section_label += (
                f" (+{branch_count} branch{'es' if branch_count != 1 else ''})"
            )
        section_label += f" — {ts_to_local_str(create_time)[:10]}"
        toc.append(f"  {section_label}\n")

        # estimate lines per conversation: ~20–30 lines per branch, messages vary
        msg_count = len(root.get("msgs", []))
        est_lines = max(10, msg_count // 3)
        for _ in items[1:]:
            est_lines += max(8, len(_.get("msgs", [])) // 5)
        line_num += est_lines + 3  # 3 for separators

    toc.append("\n")
    return toc


def _build_clean_txt(
    group_order: List[Tuple[str, List[Dict[str, Any]]]],
    topics: List[str],
    root: Path,
) -> str:
    """Build a clean, AI-friendly TXT format without technical noise.

    Removes:
      - Inline timestamps and IDs (kept minimal in metadata block)
      - Branch/root technical labels (replaced with clear section headers)
      - Excessive markdown remnants

    Includes:
      - Clear conversation grouping (thread name at top, branches below)
      - Message separation (readable dashes)
      - Sources registry at end
      - Minimal metadata header
    """
    out: List[str] = []
    topic_label = ", ".join(topics) if topics else "Dossier"

    # === HEADER ===
    out.append(f"DOSSIER: {topic_label}\n")
    out.append(
        f"Generated: {ts_to_local_str(datetime.now(tz=timezone.utc).timestamp())}\n"
    )
    out.append(f"Source: {root}\n")
    out.append("\n")

    # === TOC ===
    # Convert group_order back to dict format for TOC generation
    groups_dict = dict(group_order)
    toc = _generate_toc(groups_dict)
    out.extend(toc)

    # === CONVERSATIONS ===
    conv_num = 0
    all_sources: List[Tuple[str, str]] = []

    for _, items in group_order:
        conv_num += 1
        root_item = items[0]
        root_title = root_item["title"] or "Untitled"
        root_msgs = root_item["msgs"]

        # Section header (clean, no technical IDs unless needed for reference)
        out.append(f"\n{'='*70}\n")
        out.append(f"{conv_num}. {root_title}\n")
        out.append(f"{'='*70}\n\n")

        # Root conversation messages
        if root_msgs:
            for msg in root_msgs:
                role = msg.role.capitalize()
                out.append(f"{role}:\n\n{msg.text}\n\n")
                # extract sources incrementally
                all_sources.extend(_extract_sources(msg.text))
        else:
            out.append("[No messages in root conversation.]\n\n")

        # Branches
        for branch_idx, branch_item in enumerate(items[1:], start=1):
            branch_title = branch_item["title"] or "Untitled"
            branch_msgs = branch_item["msgs"]

            out.append(f"\n--- Branch {branch_idx}: {branch_title} ---\n\n")

            if branch_msgs:
                for msg in branch_msgs:
                    role = msg.role.capitalize()
                    out.append(f"{role}:\n\n{msg.text}\n\n")
                    all_sources.extend(_extract_sources(msg.text))
            else:
                out.append("[No new messages in this branch.]\n\n")

    # === SOURCES REGISTRY ===
    if all_sources:
        # Deduplicate and sort
        sources_dict: Dict[str, str] = {}
        for url, label in all_sources:
            if url not in sources_dict:
                sources_dict[url] = label

        out.append(f"\n{'='*70}\n")
        out.append("SOURCES REGISTRY\n")
        out.append(f"{'='*70}\n\n")
        for i, (url, label) in enumerate(sorted(sources_dict.items()), start=1):
            out.append(f"[{i}] {label}\n    {url}\n\n")

    return "".join(out)


def _strip_tool_noise(text: str) -> str:
    """Remove tool-call JSON blocks and other technical noise.

    Removes patterns like:
      {"search_query": "..."}
      {"task_violates_safety_guidelines": ...}
      {"open": ...}
      Tool creation/update messages
      Meta-prompt instructions
    """
    # Extended tool-related keys (including system/safety fields)
    tool_keys = [
        "search_query",
        "tool_call",
        "function",
        "action",
        "open",
        "click",
        "find",
        "screenshot",
        "response_length",
        "file_path",
        "command",
        "terminal",
        "browser",
        "task_violates_safety_guidelines",
        "updates",
        "comments",
        "title",
        "prompt",
    ]

    # Remove complete JSON objects that contain tool-related keys
    pattern = r"\{\s*['\"]?(?:" + "|".join(tool_keys) + r")['\"]?.*?\}"
    text = re.sub(pattern, "", text, flags=re.DOTALL)

    # Remove tool call markers
    text = re.sub(r"\[tool_call:.*?\]", "", text, flags=re.DOTALL)

    # Remove lines starting with 'tool' or '**tool**'
    text = re.sub(r"^\s*(?:\*\*)?tool\s*\(.*", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*(?:\*\*)?tool\s+\w+.*", "", text, flags=re.MULTILINE)

    # Remove tool creation/update messages
    text = re.sub(
        r"(?:Successfully created|Successfully updated|Failed with error).*?\n",
        "",
        text,
    )

    # Remove meta-prompt instructions like "Make sure to include...", "remember to..."
    text = re.sub(
        r"^[^\n]*(?:Make sure to|remember to|don't forget to).*?$",
        "",
        text,
        flags=re.MULTILINE | re.IGNORECASE,
    )

    # Remove tool usage boilerplate/instruction blocks that leak into transcripts
    text = re.sub(
        r"^##\s+How to invoke the file_search tool.*?(?=^##\s|\Z)",
        "",
        text,
        flags=re.DOTALL | re.MULTILINE | re.IGNORECASE,
    )
    text = re.sub(
        r"^##\s+How to handle results from file_search.*?(?=^##\s|\Z)",
        "",
        text,
        flags=re.DOTALL | re.MULTILINE | re.IGNORECASE,
    )
    text = re.sub(
        r"^##\s+Tool usage instructions.*?(?=^##\s|\Z)",
        "",
        text,
        flags=re.DOTALL | re.MULTILINE | re.IGNORECASE,
    )
    text = re.sub(
        r"^The file is too long and its contents have been truncated\..*?$",
        "",
        text,
        flags=re.MULTILINE | re.IGNORECASE,
    )

    # Remove lines that are pure JSON (start with { and contain tool keys)
    lines = text.split("\n")
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        normalized = re.sub(r"[\u200b-\u200f\u2060\ufeff]", "", stripped)
        # Skip JSON/Tool Call annotation lines outright
        if re.search(r"\[\s*JSON\s*/\s*Tool\s*Call\s*\]", normalized, re.IGNORECASE):
            continue
        # Skip JSON blocks with tool/system keys
        if normalized.startswith("{") and any(key in normalized for key in tool_keys):
            continue
        # Skip lines with embedded JSON containing safety/task fields
        if any(
            f'"{key}"' in normalized or f"'{key}'" in normalized
            for key in ["task_violates_safety_guidelines", "updates", "comments"]
        ):
            continue
        cleaned_lines.append(line)
    text = "\n".join(cleaned_lines)

    # Remove excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _strip_citation_markers(text: str) -> str:
    """Remove internal citation markers and navigation lists.

    Removes:
      - <citeturn0newsXX> tags
      - <navlist>...</navlist> blocks
      - Bare turn0newsXX tokens
      - Citation reference artifacts
      - Internal phantom citation markers: 【…†L…】
    """
    # Remove <citeturn...> tags
    text = re.sub(r"<citeturn[^>]*>", "", text, flags=re.IGNORECASE)

    # Remove <navlist>...</navlist> blocks
    text = re.sub(r"<navlist>.*?</navlist>", "", text, flags=re.DOTALL | re.IGNORECASE)

    # Remove bare turn0newsXX tokens
    text = re.sub(r"\bturn\d+news\d+\b", "", text, flags=re.IGNORECASE)

    # Remove citation brackets like [1], [2] that appear alone
    text = re.sub(r"\[\d+\](?!\S)", "", text)

    # Remove internal phantom citation markers: 【…†L…】 or similar CJK bracket patterns
    text = re.sub(r"【[^】]*†[^】]*】", "", text)
    text = re.sub(r"【[^】]*】", "", text)  # Remove other 【】 blocks

    # Remove residual ref placeholders
    text = re.sub(r"\[REF REMOVED\]", "", text, flags=re.IGNORECASE)

    # Clean up excessive whitespace from removals
    text = re.sub(r" {2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def _sanitize_openai_markup(text: str) -> str:
    """Remove OpenAI/ChatGPT UI markup blobs from output.

    Removes:
      - <img.../> tags (image placeholders)
      - <click.../> tags (interactive elements)
      - <link.../> tags
      - <audio.../> tags
      - Other angle-bracket UI markup (generic)
      - Stray appendix headers (to prevent duplication)
    """
    # Remove self-closing angle-bracket tags: <tag.../> or <tag>
    text = re.sub(r"</?[a-zA-Z][^>]*/?>\s*", "", text)

    # Remove any remaining OpenAI-style span markers (often empty or metadata)
    text = re.sub(r"<span[^>]*>.*?</span>", "", text, flags=re.DOTALL)

    # Remove stray "APPENDIX" headers to prevent duplication
    # (the real header will be emitted once by the pipeline)
    text = re.sub(
        r"APPENDIX:\s*RESEARCH LOG & TOOL ARTIFACTS\s*",
        "",
        text,
        flags=re.IGNORECASE,
    )

    # Remove private-use markers and spans bounded by them
    text = re.sub(r"[\ue000-\uf8ff].*?[\ue000-\uf8ff]", "", text, flags=re.DOTALL)
    text = re.sub(r"[\ue000-\uf8ff]", "", text)

    # Remove common non-printing artifacts (replacement/noncharacter/soft hyphen)
    text = re.sub(r"[\u00ad\ufffd\ufffe]", "", text)

    # Remove turn/citeturn tokens if any remain
    text = re.sub(
        r"(?:cite)?turn\d+(?:search|news|view|file)\w*",
        "",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"\bciteturn\d+\w+\b", "", text, flags=re.IGNORECASE)

    # Remove separator lines often associated with appendix (lines of = or -)
    text = re.sub(r"^={60,}.*?$", "", text, flags=re.MULTILINE)

    # Clean up whitespace damage from removals
    text = re.sub(r" {2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def _is_appendix_header_line(line: str) -> bool:
    """Detect appendix header lines robustly (ignores punctuation/soft hyphens)."""
    normalized = re.sub(r"[^A-Za-z]", "", line).upper()
    return "APPENDIX" in normalized and "RESEARCHLOGTOOLARTIFACTS" in normalized


def _strip_existing_appendix(text: str) -> str:
    """Remove any existing appendix block to prevent duplication."""
    lines = text.split("\n")
    for idx, line in enumerate(lines):
        if _is_appendix_header_line(line):
            return "\n".join(lines[:idx]).rstrip()
    return text


def _dedupe_appendix_header(text: str) -> str:
    """Ensure appendix header appears only once, preserving content."""
    marker = "APPENDIX: RESEARCH LOG & TOOL ARTIFACTS"
    if marker not in text:
        return text
    parts = text.split(marker)
    if len(parts) <= 2:
        return text
    # Keep first marker and remove subsequent marker occurrences
    first = parts[0] + marker + parts[1]
    rest = "".join(parts[2:]).replace(marker, "")
    return first + rest


def _remove_appendix_header_lines(text: str) -> str:
    """Remove any appendix header lines from text."""
    lines = [line for line in text.split("\n") if not _is_appendix_header_line(line)]
    return "\n".join(lines)


def _replace_dead_citations(text: str, sources_dict: Dict[str, str]) -> str:
    """Replace dead citation tokens (turn0news01, etc.) with source registry pointers.

    If a citation token matches a source, replace with [SOURCE: {index}].
    Otherwise, replace with [SOURCE: see registry] as a placeholder.

    NOTE: This is now largely handled by _strip_citation_markers().
    Kept for backward compatibility.
    """
    # Most work done by _strip_citation_markers now
    return text


def _tag_sources(
    sources: List[Tuple[str, str]], used_links: Optional[Set[str]] = None
) -> Dict[str, List[Tuple[str, str]]]:
    """Tag sources by category and usage status.

    Returns Dict[category, List[(url, label)]]:
      - used_in_drafts: Links from used_links set (highest priority)
      - candidate: Links with keywords suggesting relevance for next column
      - legal: .gov.br, .senado, .camara, legal-related domains
      - media: news, .com.br news outlets, press
      - economic: economic, financial, trade, central bank
      - internal: notes, transcripts, internal documents
      - other: everything else (lowest priority)
    """
    categories = {
        "used_in_drafts": [],
        "candidate": [],
        "legal": [],
        "media": [],
        "economic": [],
        "internal": [],
        "other": [],
    }

    # Keywords for identifying candidates (generic research-relevant terms)
    candidate_keywords = [
        "research",
        "analysis",
        "report",
        "study",
        "project",
        "draft",
        "document",
        "paper",
        "article",
        "summary",
    ]

    legal_keywords = [
        ".gov.br",
        ".senado",
        ".camara",
        "judicial",
        "legal",
        "court",
        "law",
    ]
    media_keywords = [
        "news",
        "folha",
        "globo",
        "estadao",
        "uol",
        "bbc",
        "cnn",
        "press",
        "media",
        "jornalismo",
        "nytimes",
        "wsj",
        "reuters",
    ]
    economic_keywords = [
        "economic",
        "financial",
        "trade",
        "economy",
        "banco",
        "bcb",
        "imf",
        "world bank",
        "commerce",
        "bloomberg",
        "forbes",
    ]
    internal_keywords = ["note", "transcript", "internal", "memo", "meeting", "summary"]

    for url, label in sources:
        url_lower = url.lower()
        label_lower = label.lower()

        # First priority: check if in used_links
        if used_links and url in used_links:
            categories["used_in_drafts"].append((url, label))
            continue

        # Second priority: check if candidate for next column
        if any(kw in url_lower or kw in label_lower for kw in candidate_keywords):
            categories["candidate"].append((url, label))
            continue

        # Then categorize by domain
        categorized = False
        if any(kw in url_lower or kw in label_lower for kw in legal_keywords):
            categories["legal"].append((url, label))
            categorized = True
        elif any(kw in url_lower or kw in label_lower for kw in media_keywords):
            categories["media"].append((url, label))
            categorized = True
        elif any(kw in url_lower or kw in label_lower for kw in economic_keywords):
            categories["economic"].append((url, label))
            categorized = True
        elif any(kw in url_lower or kw in label_lower for kw in internal_keywords):
            categories["internal"].append((url, label))
            categorized = True

        # Default to other
        if not categorized:
            categories["other"].append((url, label))

    return categories


def _deduplicate_blocks(text: str, min_block_size: int = 200) -> str:
    """Remove duplicate blocks of text (e.g., repeated transcripts).

    Uses hash-based deduplication. Splits text by paragraphs and removes
    consecutive blocks that are similar (hash match).
    """
    paragraphs = text.split("\n\n")
    seen_hashes = {}
    deduplicated = []

    for para in paragraphs:
        if len(para) < min_block_size:
            # Keep small blocks (headers, short messages)
            deduplicated.append(para)
        else:
            # Hash larger blocks for deduplication
            para_hash = hash(normalize_text(para))
            if para_hash not in seen_hashes:
                seen_hashes[para_hash] = True
                deduplicated.append(para)
            # else: skip duplicate

    return "\n\n".join(deduplicated)


def _extract_deliverables(text: str, patterns: Optional[List[str]] = None) -> str:
    """Extract deliverables only (headings, constraints, drafts).

    If patterns provided, match against section headers.
    Default patterns: "##", "Constraint", "Draft", "Decision", "Output", "Result"
    """
    if patterns is None:
        patterns = [
            "##",
            "constraint",
            "draft",
            "decision",
            "output",
            "result",
            "deliverable",
        ]

    lines = text.split("\n")
    deliverables = []
    in_section = False
    current_section = []

    for line in lines:
        line_lower = line.lower()

        # Check if line matches any pattern (start of a deliverable)
        matches_pattern = any(p.lower() in line_lower for p in patterns)

        if matches_pattern:
            # Save previous section if any
            if current_section:
                deliverables.extend(current_section)
            # Start new section
            in_section = True
            current_section = [line]
        elif in_section:
            # Continue collecting section content
            if line.strip():  # Include non-empty lines
                current_section.append(line)
            elif current_section and len(current_section) > 1:
                # Stop section on blank line (after content)
                deliverables.extend(current_section)
                current_section = []
                in_section = False

    # Include final section
    if current_section:
        deliverables.extend(current_section)

    return "\n".join(deliverables)


def _generate_working_index(
    text: str,
    conversations: Optional[List[Dict[str, Any]]] = None,
    topics: Optional[List[str]] = None,
) -> List[str]:
    """Auto-generate navigational index with timeline and priority threads.

    Includes:
      - Global timeline (conversations by date)
      - Priority threads (top 5 by recency + keywords)
      - Section headers for navigation
    """
    index_lines = ["## WORKING INDEX\n\n"]

    # Add global timeline if conversations provided
    if conversations:
        index_lines.append("### Timeline\n\n")
        # Sort conversations by create_time
        sorted_convs = sorted(
            conversations, key=lambda c: c.get("create_time", 0), reverse=True
        )
        for conv in sorted_convs[:10]:  # Show latest 10
            cid, title = conv_id_and_title(conv)
            ctime = conv.get("create_time", 0)
            date_str = ts_to_local_str(ctime).split()[0] if ctime else "Unknown"
            index_lines.append(f"  - {date_str}: {title} (ID: {cid[:8]}...)\n")
        index_lines.append("\n")

    # Add priority threads (based on keywords and recency)
    if conversations and topics:
        index_lines.append("### Priority Threads (Read These First)\n\n")
        # Generic priority keywords (project/deliverable focused)
        priority_keywords = [
            "draft",
            "decision",
            "deliverable",
            "output",
            "final",
            "review",
            "summary",
            "analysis",
        ]

        scored_convs = []
        for conv in conversations:
            cid, title = conv_id_and_title(conv)
            score = 0
            title_lower = (title or "").lower()

            # Score by keyword presence
            for kw in priority_keywords:
                if kw in title_lower:
                    score += 2

            # Score by topic match
            for topic in topics:
                if topic.lower() in title_lower:
                    score += 3

            # Boost recent conversations
            ctime = conv.get("create_time", 0)
            if ctime:
                # Conversations from last 30 days get boost
                import time

                days_ago = (time.time() - ctime) / 86400
                if days_ago < 30:
                    score += (30 - days_ago) / 10

            if score > 0:
                scored_convs.append((score, cid, title, ctime))

        # Sort by score and show top 5
        scored_convs.sort(reverse=True, key=lambda x: x[0])
        for i, (score, cid, title, ctime) in enumerate(scored_convs[:5], 1):
            date_str = ts_to_local_str(ctime).split()[0] if ctime else "Unknown"
            index_lines.append(f"  {i}. [{date_str}] {title}\n")
        index_lines.append("\n")

    # Add section navigation
    index_lines.append("### Sections\n\n")
    section_num = 0

    # Find all section headers (##, ===, etc.)
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if line.startswith("##") or line.startswith("===") or re.match(r"^\d+\.", line):
            section_num += 1
            header = line.strip().lstrip("#").strip().lstrip("=").strip()
            if header and header != "WORKING INDEX":  # Don't index ourselves
                index_lines.append(f"  {section_num:02d}. Line ~{i}: {header}\n")

    index_lines.append("\n---\n\n")
    return index_lines


def _reorganize_sources_section(
    text: str, used_links: Optional[Set[str]] = None
) -> str:
    """Extract sources from text and reorganize them by category."""
    sources_match = re.search(
        r"^={70}\nSOURCES REGISTRY\n={70}\n\n(.+)$",
        text,
        flags=re.MULTILINE | re.DOTALL,
    )
    if not sources_match:
        return text

    sources: List[Tuple[str, str]] = []
    source_pattern = r"\[(\d+)\]\s+(.+?)\n\s+(https?://\S+)"
    for match in re.finditer(source_pattern, sources_match.group(1)):
        label, url = match.group(2).strip(), match.group(3).strip()
        sources.append((url, label))

    if not sources:
        return text

    categorized = _tag_sources(sources, used_links)
    new_section = ["=" * 70 + "\n", "SOURCES REGISTRY\n", "=" * 70 + "\n\n"]

    display_order = [
        ("used_in_drafts", "**Used in Drafts**"),
        ("candidate", "**Candidates for Next Column**"),
        ("legal", "Legal Sources"),
        ("media", "Media Sources"),
        ("economic", "Economic Sources"),
        ("internal", "Internal Documents"),
        ("other", "Other Sources"),
    ]

    num = 1
    for key, label in display_order:
        if categorized.get(key):
            new_section.append(f"{label}:\n\n")
            for url, lbl in sorted(categorized[key], key=lambda x: x[1].lower()):
                new_section.append(f"[{num}] {lbl}\n    {url}\n\n")
                num += 1
            new_section.append("\n")

    return text[: sources_match.start()] + "".join(new_section)


# ===== CONFIG-DRIVEN FILTERING & SCORING =====


def load_column_config(config_file: str) -> Dict[str, Any]:
    """Load column-specific constraints from JSON config file."""
    config_path = Path(config_file).expanduser().resolve()
    if not config_path.exists():
        die(f"Config file not found: {config_file}")
    try:
        with open(config_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        die(f"Error loading config: {e}")


def matches_thread_filter(
    title: str, config: Dict[str, Any]
) -> Tuple[bool, Optional[str]]:
    """Check if thread title matches include/exclude filters. Returns (include, tag)."""
    if not title:
        return False, None

    title_lower = title.lower()
    filters = config.get("thread_filters", {})
    exclude_list = filters.get("exclude", [])
    include_dict = filters.get("include", {})

    # Check exclude list first
    for exclude_term in exclude_list:
        if exclude_term.lower() in title_lower:
            return False, None

    # Check include buckets
    for bucket_name, terms in include_dict.items():
        for term in terms:
            if term.lower() in title_lower:
                return True, bucket_name

    return False, None


def score_segment(text: str, config: Dict[str, Any]) -> float:
    """Score a text segment for relevance based on mechanism/bridging terms."""
    if not text or len(text) < 50:
        return 0.0

    scoring = config.get("segment_scoring", {})
    mechanism_terms = scoring.get("mechanism_terms", [])
    bridging_terms = scoring.get("bridging_terms", [])

    text_lower = text.lower()
    score = 0.0

    # Mechanism terms (higher weight)
    for term in mechanism_terms:
        if term.lower() in text_lower:
            score += 2.0

    # Bridging terms (lower weight)
    for term in bridging_terms:
        if term.lower() in text_lower:
            score += 1.0

    # Normalize by text length (avoid inflating on longer texts)
    normalized = score / (len(text) / 100.0) if text else 0.0
    return min(normalized, 10.0)  # Cap at 10


def extract_segments_with_context(
    msgs: List[Msg], config: Dict[str, Any], min_score: float = 0.5
) -> List[Msg]:
    """Extract only high-scoring segments + context windows."""
    if not msgs:
        return []

    context_window = config.get("segment_scoring", {}).get("context_window", 2)

    # Find high-scoring messages
    scored = []
    for i, msg in enumerate(msgs):
        score = score_segment(msg.text, config)
        scored.append((i, msg, score))

    # Keep only high-scoring
    high_scoring_indices = {i for i, m, s in scored if s >= min_score}

    if not high_scoring_indices:
        return msgs  # Fallback: return all if nothing scores high

    # Add context around high-scoring messages
    to_keep = set()
    for idx in high_scoring_indices:
        for j in range(
            max(0, idx - context_window), min(len(msgs), idx + context_window + 1)
        ):
            to_keep.add(j)

    return [msgs[i] for i in sorted(to_keep)]


def get_thread_tag(title: str, config: Dict[str, Any]) -> str:
    """Get tag for thread based on bucket from config filter."""
    _, bucket = matches_thread_filter(title, config)
    if not bucket:
        return "OTHER"

    # Generate tag from bucket name (e.g., "primary_research" -> "PRIMARY")
    # Use first word of bucket, uppercase, max 10 chars
    tag = bucket.split("_")[0].upper()[:10]
    return tag if tag else "OTHER"


def generate_completeness_check(
    convs: List[Dict[str, Any]], config: Dict[str, Any]
) -> str:
    """Generate completeness metadata line with basic statistics."""
    if not convs:
        return "No conversations found."

    # Find date range
    dates = []
    for conv in convs:
        if conv.get("create_time"):
            dates.append(conv["create_time"])

    if not dates:
        return "No date information available."

    latest_date = max(dates)
    latest_str = ts_to_local_str(latest_date).split()[0]

    # Count matches after a certain date
    now = datetime.now(tz=timezone.utc).timestamp()
    recent_matches = sum(1 for d in dates if now - d < 7 * 86400)

    metadata = (
        f"Searched conversations up to {ts_to_local_str(now).split()[0]}.\n"
        f"Last relevant match: {latest_str}.\n"
        f"Recent matches (< 7 days): {recent_matches}.\n"
        f"Total conversations in dossier: {len(convs)}."
    )

    return metadata


def generate_control_layer(config: Dict[str, Any]) -> str:
    """Generate front-matter control layer with scope router + constraints."""
    lines = [
        "=" * 70,
        "CONTROL LAYER — " + config.get("column_name", "Report"),
        "=" * 70,
        "",
    ]

    control_sections = config.get("control_layer_sections", {})
    op_constraints = config.get("op_v2_constraints", [])

    # Scope router
    if "scope_router" in control_sections:
        lines.append("SCOPE ROUTER")
        lines.append("")
        lines.append(control_sections["scope_router"])
        lines.append("")

    # Do-not-repeat rules
    if "do_not_repeat_rules" in control_sections:
        lines.append("DO-NOT-REPEAT RULES")
        lines.append("")
        for rule in control_sections["do_not_repeat_rules"]:
            lines.append(f"• {rule}")
        lines.append("")

    # Mechanism focus
    if "mechanism_focus" in control_sections:
        lines.append("MECHANISM FOCUS (from OP v2)")
        lines.append("")
        lines.append(control_sections["mechanism_focus"])
        lines.append("")

    # Evidence vs inference
    if "evidence_vs_inference" in control_sections:
        lines.append("EVIDENCE VS INFERENCE")
        lines.append("")
        lines.append(control_sections["evidence_vs_inference"])
        lines.append("")

    # Stress tests
    if "stress_tests" in control_sections:
        lines.append("STRESS TESTS")
        lines.append("")
        for test in control_sections["stress_tests"]:
            lines.append(f"• {test}")
        lines.append("")

    lines.append("=" * 70)
    lines.append("")

    return "\n".join(lines)


def _get_short_tag(bucket_name: Optional[str]) -> str:
    """
    Map config bucket names to short machine-readable tags.
    Derives tag dynamically from bucket name (e.g., 'primary_research' -> 'PRIMARY').
    """
    if not bucket_name:
        return "OTHER"

    # Generate tag from bucket name: use first word, uppercase, max 10 chars
    tag = bucket_name.split("_")[0].upper()[:10]
    return tag if tag else "OTHER"


def extract_research_artifacts(text: str) -> Tuple[str, List[str]]:
    """
    Extract tool-call artifacts (JSON, search fragments, etc.) into separate appendix.
    Returns (cleaned_text, artifacts_list) where artifacts_list is list of artifact strings.
    Also filters out any stray appendix headers that may have been included in text.
    """
    import re

    artifacts = []
    cleaned_lines = []

    # Patterns to detect and quarantine
    artifact_patterns = [
        (r"\[Search Query\].*?(?=\n\n|\n[A-Z]|$)", "Search Fragment"),
        (r"\[JSON/Tool Call\].*", "JSON/Tool Call"),
        (r"\{\".*?\}", "JSON/Tool Call"),
        (r"\[Image.*?\]", "Image Reference"),
        (r"\[GPT Model.*?\]", "Model Info"),
        (r"\[Citation Widget.*?\]", "Citation Widget"),
    ]

    lines = text.split("\n")
    for i, line in enumerate(lines):
        normalized_line = re.sub(r"[\u200b-\u200f\u2060\ufeff]", "", line)
        # Skip lines that are stray appendix headers
        if _is_appendix_header_line(normalized_line):
            continue

        # Always drop JSON/Tool Call lines (do not propagate to appendix)
        if re.search(
            r"\[\s*JSON\s*/\s*Tool\s*Call\s*\]", normalized_line, re.IGNORECASE
        ):
            continue

        found_artifact = False

        for pattern, label in artifact_patterns:
            match = re.search(pattern, normalized_line, re.DOTALL)
            if match:
                artifact_text = match.group(0)[:200]
                # Avoid propagating appendix header text into artifacts
                if _is_appendix_header_line(artifact_text):
                    found_artifact = True
                    break
                # Always drop JSON/Tool Call lines (do not propagate to appendix)
                if label == "JSON/Tool Call":
                    found_artifact = True
                    break
                # Only capture substantial artifacts for other labels
                if len(match.group(0)) > 20:
                    artifacts.append(f"[{label}] {artifact_text}...")
                    found_artifact = True
                    break

        if not found_artifact:
            cleaned_lines.append(line)

    cleaned_text = "\n".join(cleaned_lines)

    return cleaned_text, artifacts


def _generate_working_index_with_tags(
    txt: str,
    conversations: List[Dict[str, Any]] = None,
    topics: List[str] = None,
    config: Optional[Dict[str, Any]] = None,
) -> Tuple[List[str], List[str]]:
    """
    Generate working index with thread tags derived from config buckets.
    Returns (index_lines, coverage_report_lines).
    """
    if not conversations:
        return [], []

    index_lines = ["PRIORITY THREADS (with category tags)\n", "=" * 70 + "\n"]

    priority_threads = []
    included_count = 0
    tag_counts: Dict[str, int] = {}  # Dynamic tag counting

    for c in conversations:
        cid, title = conv_id_and_title(c)
        if cid and title:
            # Get tag from config if available
            bucket_tag = None
            if config:
                included, bucket_tag = matches_thread_filter(title, config)

            # Map bucket name to short tag
            short_tag = _get_short_tag(bucket_tag)

            ctime = float(c.get("create_time") or 0.0)
            priority_threads.append((cid, title, ctime, short_tag))
            included_count += 1
            tag_counts[short_tag] = tag_counts.get(short_tag, 0) + 1

    # Sort by creation time
    priority_threads.sort(key=lambda x: x[2])

    for cid, title, ctime, tag in priority_threads:
        # Always show tag in brackets (mandatory)
        tag_str = f"[{tag}] "
        index_lines.append(
            f"{tag_str}{title}\n" f"  ID: {cid} | Created: {ts_to_local_str(ctime)}\n\n"
        )

    # Generate coverage report with dynamic tags
    coverage_lines = [
        "\n" + "=" * 70,
        "COVERAGE AUDIT",
        "=" * 70,
        f"Included threads (total): {included_count}",
    ]
    for tag, count in sorted(tag_counts.items()):
        coverage_lines.append(f"  - [{tag}]: {count}")
    coverage_lines.append("")

    return index_lines, coverage_lines


# ===== CONFIG-DRIVEN FILTERING & SCORING END =====


# ----------------- excerpting + branching -----------------


# ----------------- excerpting + branching -----------------


def excerpt_messages(msgs: List[Msg], pattern: re.Pattern, context: int) -> List[Msg]:
    if not msgs:
        return []
    hits = [i for i, m in enumerate(msgs) if pattern.search(m.text)]
    if not hits:
        return []
    keep = set()
    for i in hits:
        for j in range(max(0, i - context), min(len(msgs), i + context + 1)):
            keep.add(j)
    return [msgs[i] for i in sorted(keep)]


def base_title(title: str) -> str:
    t = (title or "").strip()
    t = re.sub(r"^\s*Branch\s*[·\-:]\s*", "", t, flags=re.IGNORECASE)
    return t.strip()


def longest_common_prefix_len(
    a: List[Tuple[str, str]], b: List[Tuple[str, str]]
) -> int:
    n = min(len(a), len(b))
    i = 0
    while i < n and a[i] == b[i]:
        i += 1
    return i


def trim_branch_new_part(root_msgs: List[Msg], branch_msgs: List[Msg]) -> List[Msg]:
    ra = [(m.role, normalize_text(m.text)) for m in root_msgs]
    rb = [(m.role, normalize_text(m.text)) for m in branch_msgs]
    k = longest_common_prefix_len(ra, rb)
    return branch_msgs[k:]


# ----------------- dossier core -----------------


def compile_topic_pattern(topics: List[str]) -> re.Pattern:
    parts = [re.escape(t) for t in topics if t.strip()]
    if not parts:
        # never match
        return re.compile(r"a^")
    return re.compile("|".join(parts), re.IGNORECASE)


def build_combined_dossier(
    *,
    topics: List[str],
    mode: str,
    context: int,
    root: Path,
    dossiers_dir: Path,
    wanted_ids: List[str],
    convs: List[Dict[str, Any]],
    formats: Optional[List[str]] = None,
    split: bool = False,
    patterns: Optional[List[str]] = None,
    dedup: bool = True,
    used_links_file: Optional[str] = None,
    config_file: Optional[str] = None,
    name: Optional[str] = None,
) -> Path:
    if not wanted_ids:
        die("No valid selections provided; cannot build dossier.")

    by_id: Dict[str, Dict[str, Any]] = {}
    for c in convs:
        cid, _ = conv_id_and_title(c)
        if cid:
            by_id[cid] = c

    missing = [i for i in wanted_ids if i not in by_id]
    if missing:
        die("Some IDs not found in export:\n" + "\n".join(missing))

    topic_re = compile_topic_pattern(topics)
    topic_label = ", ".join(topics)

    convo_items = []
    for cid in wanted_ids:
        c = by_id[cid]
        _, title = conv_id_and_title(c)
        ctime = float(c.get("create_time") or 0.0)
        msgs = extract_messages_best_effort(c)
        convo_items.append(
            {
                "id": cid,
                "title": title,
                "base_title": base_title(title),
                "ctime": ctime,
                "msgs": msgs,
            }
        )

    groups: Dict[str, List[Dict[str, Any]]] = {}
    for item in convo_items:
        groups.setdefault(item["base_title"] or item["title"] or item["id"], []).append(
            item
        )

    for k in groups.keys():
        groups[k].sort(key=lambda x: x["ctime"])

    group_order = sorted(
        groups.items(), key=lambda kv: kv[1][0]["ctime"] if kv[1] else 0.0
    )

    # Determine output directory and filename
    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    if name:
        # Create named subfolder: dossiers/{name}/
        named_dir = dossiers_dir / safe_slug(name)
        named_dir.mkdir(parents=True, exist_ok=True)
        out_name = f"{timestamp}.md"
        out_path = named_dir / out_name
    else:
        # Flat structure (original behavior): dossiers/dossier__{topic}__{timestamp}.md
        out_name = f"dossier__{safe_slug(topic_label)}__{timestamp.replace('-', '')}.md"
        out_path = dossiers_dir / out_name

    doc: List[str] = []
    doc.append(f"# Dossier: {topic_label}\n\n")
    doc.append(
        f"- generated_at: {ts_to_local_str(datetime.now(tz=timezone.utc).timestamp())}\n"
    )
    doc.append(f"- export_root: {root}\n")
    doc.append(f"- mode: {mode}\n\n")
    doc.append("---\n\n")

    for _, items in group_order:
        root_item = items[0]
        root_id = root_item["id"]
        root_title = root_item["title"] or "Untitled"
        doc.append(f"## Thread: {root_title}\n\n")
        doc.append(f"- root_id: {root_id}\n")
        doc.append(
            f"- conversation_create_time: {ts_to_local_str(root_item['ctime'])}\n\n"
        )

        root_msgs = root_item["msgs"]
        if mode == "excerpts":
            root_msgs = excerpt_messages(root_msgs, topic_re, context)

        if root_msgs:
            doc.append("### Root conversation\n\n")
            for m in root_msgs:
                doc.append(f"**{m.role}** ({ts_to_local_str(m.t)})\n\n{m.text}\n\n")
        else:
            doc.append(
                "_No matching excerpts in root conversation._\n\n"
                if mode == "excerpts"
                else "_No messages found._\n\n"
            )

        for b in items[1:]:
            b_id = b["id"]
            b_title = b["title"] or "Untitled"
            b_msgs_new = trim_branch_new_part(root_item["msgs"], b["msgs"])

            if mode == "excerpts":
                b_msgs_new = excerpt_messages(b_msgs_new, topic_re, context)

            doc.append(f"### Branch: {b_title}\n\n")
            doc.append(f"- branch_id: {b_id}\n")
            doc.append(
                f"- branch_conversation_create_time: {ts_to_local_str(b['ctime'])}\n\n"
            )

            if b_msgs_new:
                for m in b_msgs_new:
                    doc.append(f"**{m.role}** ({ts_to_local_str(m.t)})\n\n{m.text}\n\n")
            else:
                doc.append(
                    "_No matching excerpts in this branch._\n\n"
                    if mode == "excerpts"
                    else "_No new messages after trimming._\n\n"
                )

        doc.append("---\n\n")

    md_content = "".join(doc)
    # Normalize requested formats: default to ['txt'] when not provided
    req_formats = [f.lower() for f in (formats or [])]
    if not req_formats:
        req_formats = ["txt"]

    # Utility to convert markdown -> plain text
    def _markdown_to_plain(md: str) -> str:
        s = re.sub(r"\r\n?", "\n", md)
        s = re.sub(r"(?m)^#{1,6}\s*", "", s)
        s = re.sub(r"\*\*(.*?)\*\*", r"\1", s)
        s = re.sub(r"\*(.*?)\*", r"\1", s)
        s = re.sub(r"```.*?```", "", s, flags=re.S)
        s = re.sub(r"`([^`]+)`", r"\1", s)
        s = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", s)
        s = re.sub(r"\n{3,}", "\n\n", s)
        return s.strip() + "\n"

    md_path = out_path
    created_primary: Optional[Path] = None

    # Build raw TXT (clean format, minimal processing)
    raw_txt = None
    if "txt" in req_formats:
        try:
            raw_txt = _build_clean_txt(group_order, topics, root)
        except Exception as e:
            print(f"WARNING: Raw TXT generation failed: {e}", file=sys.stderr)

    # Build working TXT if split mode enabled (apply all filters)
    working_txt = None
    append_expected = False
    if split and raw_txt:
        try:
            # Load config if provided
            config = None
            if config_file:
                config = load_column_config(config_file)

            # Load used links if file provided
            used_links: Optional[Set[str]] = None
            if used_links_file and Path(used_links_file).exists():
                try:
                    with open(used_links_file, "r", encoding="utf-8") as f:
                        used_links = {
                            line.strip()
                            for line in f
                            if line.strip() and not line.startswith("#")
                        }
                except Exception as e:
                    print(
                        f"WARNING: Could not load used_links_file: {e}", file=sys.stderr
                    )

            working_txt = raw_txt
            # Apply all processing filters (in order)
            working_txt = _strip_tool_noise(working_txt)
            working_txt = _strip_citation_markers(working_txt)
            working_txt = _sanitize_openai_markup(working_txt)
            working_txt = _strip_existing_appendix(working_txt)
            working_txt = _remove_appendix_header_lines(working_txt)
            working_txt = _replace_dead_citations(working_txt, {})
            if dedup:
                working_txt = _deduplicate_blocks(working_txt, min_block_size=200)
            if patterns is not None:
                working_txt = _extract_deliverables(working_txt, patterns)

            # Reorganize sources section with categorization
            working_txt = _reorganize_sources_section(working_txt, used_links)

            # Extract and quarantine research artifacts to list (not to string yet)
            working_txt, artifacts_list = extract_research_artifacts(working_txt)
            if artifacts_list:
                artifacts_list = [
                    a
                    for a in artifacts_list
                    if "APPENDIX: RESEARCH LOG" not in a
                    and "RESEARCH LOG & TOOL ARTIFACTS" not in a
                    and "JSON/Tool Call" not in a
                ]
            append_expected = bool(artifacts_list)

            # Add control layer front-matter if config provided
            if config:
                control_layer = generate_control_layer(config)
                completeness = generate_completeness_check(convs, config)
                front_matter = (
                    control_layer
                    + "COMPLETENESS CHECK\n"
                    + "=" * 70
                    + "\n"
                    + completeness
                    + "\n"
                    + "=" * 70
                    + "\n\n"
                )
                working_txt = front_matter + working_txt

            if not working_txt.strip():
                die("NO INCLUDED THREADS/SEGMENTS — CHECK FILTERS/KEYWORDS/PATTERNS")

            # Add working index at top (with timeline, priority threads, and tags if config)
            selected_convs = [by_id[cid] for cid in wanted_ids if cid in by_id]
            coverage_report = []
            if config:
                # Use tagged index if config provided
                working_idx, coverage_report = _generate_working_index_with_tags(
                    working_txt,
                    conversations=selected_convs,
                    topics=topics,
                    config=config,
                )
            else:
                # Fall back to standard index
                working_idx = _generate_working_index(
                    working_txt, conversations=selected_convs, topics=topics
                )

            # Build final output: index + coverage + content + appendix (once, at end)
            final_txt = "".join(working_idx)
            if coverage_report:
                final_txt += "\n" + "\n".join(coverage_report)
            final_txt += "\n" + working_txt

            # Emit appendix ONCE at the very end if artifacts exist
            if artifacts_list:
                final_txt += (
                    "\n\n" + "=" * 70 + "\n"
                    "APPENDIX: RESEARCH LOG & TOOL ARTIFACTS\n"
                    "=" * 70 + "\n\n"
                    "This section contains metadata, tool-call fragments, and provenance\n"
                    "information from the research extraction process.\n\n"
                    + "\n\n".join(artifacts_list)
                )

            # Final safety: ensure appendix header appears only once
            final_txt = _dedupe_appendix_header(final_txt)

            working_txt = final_txt
        except Exception as e:
            print(f"WARNING: Working TXT processing failed: {e}", file=sys.stderr)
            working_txt = None

    # Write MD if requested
    if "md" in req_formats:
        try:
            md_path.write_text(md_content, encoding="utf-8")
            created_primary = md_path
        except Exception as e:
            print(f"WARNING: MD generation failed: {e}", file=sys.stderr)

    # Write TXT files
    if "txt" in req_formats:
        try:
            txt_path = out_path.with_suffix(".txt")
            if raw_txt:
                txt_path.write_text(raw_txt, encoding="utf-8")
                if created_primary is None:
                    created_primary = txt_path

            # Write working variant if split enabled
            if split and working_txt:
                working_path = txt_path.parent / (
                    txt_path.stem + "__working" + txt_path.suffix
                )

                # HARD GUARDS: Verify appendix header appears exactly once (only if appended)
                if append_expected:
                    appendix_header_count = working_txt.count(
                        "APPENDIX: RESEARCH LOG & TOOL ARTIFACTS"
                    )
                    if appendix_header_count != 1:
                        print(
                            f"WARNING: Appendix header appears {appendix_header_count} times (expected 1)",
                            file=sys.stderr,
                        )

                    research_log_count = working_txt.count(
                        "RESEARCH LOG & TOOL ARTIFACTS"
                    )
                    if research_log_count != 1:
                        print(
                            f"WARNING: 'RESEARCH LOG & TOOL ARTIFACTS' appears {research_log_count} times (expected 1)",
                            file=sys.stderr,
                        )

                working_path.write_text(working_txt, encoding="utf-8")
        except Exception as e:
            print(f"WARNING: TXT generation failed: {e}", file=sys.stderr)

    # Write DOCX (use MD as source if available, fallback to plain conversion)
    if "docx" in req_formats:
        try:
            from docx import Document  # type: ignore

            docx_path = out_path.with_suffix(".docx")
            docx_doc = Document()
            plain = _markdown_to_plain(md_content)
            for para in [p for p in plain.split("\n\n") if p.strip()]:
                docx_doc.add_paragraph(para)
            docx_doc.save(str(docx_path))
            if created_primary is None:
                created_primary = docx_path
        except Exception as e:
            print(f"WARNING: DOCX generation failed: {e}", file=sys.stderr)

    if created_primary is None:
        die(
            "No dossier output files were created. "
            "Check requested formats and dependencies (for DOCX, install python-docx)."
        )
    return created_primary


# ----------------- commands -----------------


def cmd_init(args: argparse.Namespace) -> None:
    home = home_dir(args.home)
    created, existing = init_layout(home)

    quiet = bool(getattr(args, "quiet", False))
    if quiet:
        return

    print(f"Home: {home}")
    for d in created:
        print(f"created: {d.name}/")
    for d in existing:
        print(f"exists: {d.name}/")
    if not created:
        print("All required folders already exist.")


def cmd_latest_zip(args: argparse.Namespace) -> None:
    home = home_dir(args.home)
    zips_dir, _, _ = ensure_layout(home)
    print(newest_zip(zips_dir))


def cmd_extract(args: argparse.Namespace) -> None:
    home = home_dir(args.home)
    zips_dir, extracted_dir, _ = ensure_layout(home)
    zpath = Path(args.zip).expanduser().resolve() if args.zip else newest_zip(zips_dir)
    if not zpath.exists():
        die(f"ZIP not found: {zpath}")
    out_dir = extracted_dir / zpath.stem
    out_dir.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(zpath, "r") as zf:
        zf.extractall(out_dir)
    refresh_latest_symlink(extracted_dir, out_dir)
    # Print output unless quiet requested (top-level or subcommand)
    quiet = bool(getattr(args, "quiet", False))
    if not quiet:
        print(out_dir)

    # Indexing: by default update the global index at extracted/cgpt_index.db (Option A)
    no_index = bool(getattr(args, "no_index", False))
    reindex = bool(getattr(args, "reindex", False))
    if not no_index:
        try:
            db_path = extracted_dir / "cgpt_index.db"
            # If reindex, perform a full rebuild; otherwise perform incremental upsert
            index_export(out_dir, db_path, reindex=reindex, show_progress=not quiet)
        except Exception as e:
            print(f"WARNING: indexing failed: {e}", file=sys.stderr)


def cmd_ids(args: argparse.Namespace) -> None:
    home = home_dir(args.home)
    _, extracted_dir, _ = ensure_layout(home)
    root = (
        Path(args.root).expanduser().resolve()
        if args.root
        else default_root(extracted_dir)
    )
    data_file = find_conversations_json(root)
    if not data_file:
        die(f"No JSON found under {root}")
    data = load_json(data_file)
    convs = normalize_conversations(data)
    for c in convs:
        cid, title = conv_id_and_title(c)
        if cid:
            print(f"{cid}\t{title}")


def cmd_find(args: argparse.Namespace) -> None:
    query_raw = args.query.strip()
    q = query_raw.lower()
    if not query_raw:
        die("Query cannot be empty.")
    home = home_dir(args.home)
    _, extracted_dir, _ = ensure_layout(home)
    root = (
        Path(args.root).expanduser().resolve()
        if args.root
        else default_root(extracted_dir)
    )
    data_file = find_conversations_json(root)
    if not data_file:
        die(f"No JSON found under {root}")
    data = load_json(data_file)
    convs = normalize_conversations(data)
    pat = re.compile(re.escape(query_raw), re.IGNORECASE)
    for c in convs:
        cid, title = conv_id_and_title(c)
        if cid and q in (title or "").lower():
            colored = _colorize_title_with_topic(title or "", query_raw)
            print(f"{cid}\t{colored}")


def cmd_search(args: argparse.Namespace) -> None:
    # Support both legacy single positional `query` and new multi-term `--terms`.
    query_raw = getattr(args, "query", None) or ""
    terms: List[str] = []
    if getattr(args, "terms", None):
        terms = [t for t in args.terms if t]
    elif query_raw:
        terms = [query_raw]
    if not terms:
        die("Query cannot be empty.")
    and_terms = bool(getattr(args, "and_terms", False))

    # Resolve where: prefer explicit --where (`where_opt`), then positional `where`, then default 'title'
    where = getattr(args, "where_opt", None) or getattr(args, "where", None) or "title"

    home = home_dir(getattr(args, "home", None))
    _, extracted_dir, _ = ensure_layout(home)
    root = (
        Path(args.root).expanduser().resolve()
        if getattr(args, "root", None)
        else default_root(extracted_dir)
    )

    # Try using SQLite FTS index if present for faster searches
    db_path = extracted_dir / "cgpt_index.db"
    if db_path.exists():
        fts_q = build_fts_query(terms, and_terms)
        if fts_q:
            rows = query_index(db_path, fts_q, where=where)
            if rows:
                for cid, title in rows:
                    colored_title = _colorize_title_with_topics(title or "", terms)
                    print(f"{cid}\t{colored_title}")
                return

    data_file = find_conversations_json(root)
    if not data_file:
        die(f"No JSON found under {root}")
    data = load_json(data_file)
    convs = normalize_conversations(data)

    # For fallback scanning: compile an OR regex for quick checks; use per-term checks for AND.
    topic_re = compile_topic_pattern(terms)

    for c in convs:
        cid, title = conv_id_and_title(c)
        if not cid:
            continue

        hit = False
        # Title checks
        if where in ("title", "all"):
            t = title or ""
            if and_terms:
                ok_title = all((term.lower() in t.lower()) for term in terms)
            else:
                ok_title = bool(topic_re.search(t))
            if ok_title:
                hit = True

        # Message checks (AND/OR semantics)
        if not hit and where in ("messages", "all"):
            if and_terms:
                ok_msgs = True
                for term in terms:
                    p = re.compile(re.escape(term), re.IGNORECASE)
                    if not conversation_matches_text(c, p):
                        ok_msgs = False
                        break
            else:
                ok_msgs = conversation_matches_text(c, topic_re)
            if ok_msgs:
                hit = True

        if hit:
            colored_title = _colorize_title_with_topics(title or "", terms)
            print(f"{cid}\t{colored_title}")


def cmd_make_dossiers(args: argparse.Namespace) -> None:
    home = home_dir(args.home)
    _, extracted_dir, dossiers_dir = ensure_layout(home)
    root = (
        Path(args.root).expanduser().resolve()
        if args.root
        else default_root(extracted_dir)
    )
    data_file = find_conversations_json(root)
    if not data_file:
        die(f"No JSON found under {root}")
    data = load_json(data_file)
    convs = normalize_conversations(data)

    wanted: List[str] = []
    if args.ids:
        wanted.extend(args.ids)
    if args.ids_file:
        p = Path(args.ids_file).expanduser().resolve()
        if not p.exists():
            die(f"IDs file not found: {p}")
        wanted.extend(
            [
                ln.strip()
                for ln in p.read_text(encoding="utf-8").splitlines()
                if ln.strip()
            ]
        )
    wanted = [w.strip() for w in wanted if w.strip()]
    if not wanted:
        die("Provide --ids and/or --ids-file")

    by_id = {}
    for c in convs:
        cid, _ = conv_id_and_title(c)
        if cid:
            by_id[cid] = c

    missing = [i for i in wanted if i not in by_id]
    if missing:
        die("Some IDs not found in export:\n" + "\n".join(missing))

    for cid in wanted:
        c = by_id[cid]
        _, title = conv_id_and_title(c)
        base = dossiers_dir / f"{cid}__{safe_slug(title or 'untitled')}"
        md_path = base.with_suffix(".md")

        msgs = extract_messages_best_effort(c)
        header = f"# {title or 'Untitled'}\n\n- id: {cid}\n"
        ctime = c.get("create_time")
        if isinstance(ctime, (int, float)):
            header += f"- conversation_create_time: {ts_to_local_str(float(ctime))}\n"
        header += "\n---\n\n"

        parts = [header]
        for m in msgs:
            role_name = m.role.capitalize()
            parts.append(f"## {role_name} ({ts_to_local_str(m.t)})\n\n{m.text}\n\n")

        md_content = "".join(parts)

        # normalize requested formats (default to txt)
        req_formats = [f.lower() for f in (getattr(args, "format", None) or [])]
        if not req_formats:
            req_formats = ["txt"]

        # helper: convert markdown -> plain text
        def _markdown_to_plain(md: str) -> str:
            s = re.sub(r"\r\n?", "\n", md)
            s = re.sub(r"(?m)^#{1,6}\s*", "", s)
            s = re.sub(r"\*\*(.*?)\*\*", r"\1", s)
            s = re.sub(r"\*(.*?)\*", r"\1", s)
            s = re.sub(r"```.*?```", "", s, flags=re.S)
            s = re.sub(r"`([^`]+)`", r"\1", s)
            s = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", s)
            s = re.sub(r"\n{3,}", "\n\n", s)
            return s.strip() + "\n"

        created_paths: List[Path] = []

        if "md" in req_formats:
            try:
                md_path.write_text(md_content, encoding="utf-8")
                created_paths.append(md_path)
            except Exception as e:
                print(
                    f"WARNING: Failed to write Markdown file {md_path}: {e}",
                    file=sys.stderr,
                )

        if "txt" in req_formats:
            try:
                # Use clean TXT format for per-conversation dossiers as well
                clean_txt_lines: List[str] = []
                clean_txt_lines.append(f"{title or 'Untitled'}\n")
                clean_txt_lines.append(
                    f"Generated: {ts_to_local_str(datetime.now(tz=timezone.utc).timestamp())}\n"
                )
                clean_txt_lines.append(f"Source: {root}\n\n")

                # Extract sources from this conversation
                sources: List[Tuple[str, str]] = []
                for msg in msgs:
                    sources.extend(_extract_sources(msg.text))

                # Main content
                clean_txt_lines.append("=" * 70 + "\n")
                for msg in msgs:
                    role = msg.role.capitalize()
                    clean_txt_lines.append(f"{role}:\n\n{msg.text}\n\n")

                # Sources registry
                if sources:
                    sources_dict: Dict[str, str] = {}
                    for url, label in sources:
                        if url not in sources_dict:
                            sources_dict[url] = label

                    clean_txt_lines.append("\n" + "=" * 70 + "\n")
                    clean_txt_lines.append("SOURCES REGISTRY\n")
                    clean_txt_lines.append("=" * 70 + "\n\n")
                    for i, (url, label) in enumerate(
                        sorted(sources_dict.items()), start=1
                    ):
                        clean_txt_lines.append(f"[{i}] {label}\n    {url}\n\n")

                txt_path = base.with_suffix(".txt")
                txt_path.write_text("".join(clean_txt_lines), encoding="utf-8")
                created_paths.append(txt_path)
            except Exception as e:
                print(
                    f"WARNING: Failed to write TXT file for conversation {cid}: {e}",
                    file=sys.stderr,
                )

        if "docx" in req_formats:
            try:
                from docx import Document  # type: ignore

                docx_path = base.with_suffix(".docx")
                docx_doc = Document()
                plain = _markdown_to_plain(md_content)
                for para in [p for p in plain.split("\n\n") if p.strip()]:
                    docx_doc.add_paragraph(para)
                docx_doc.save(str(docx_path))
                created_paths.append(docx_path)
            except Exception as e:
                print(
                    f"WARNING: Failed to write DOCX file for conversation {cid}: {e}",
                    file=sys.stderr,
                )

        # Print whichever primary file was created (prefer txt then md then docx)
        if created_paths:
            # choose preferred ordering
            for ext in (".txt", ".md", ".docx"):
                for p in created_paths:
                    if p.suffix == ext:
                        print(p)
                        break
                else:
                    continue
                break
        else:
            print(
                f"WARNING: No output files created for conversation {cid}",
                file=sys.stderr,
            )


def cmd_build_dossier(args: argparse.Namespace) -> None:
    mode = getattr(args, "mode", None) or "full"
    context = int(args.context)

    topics: List[str] = []
    if getattr(args, "topic", None):
        topics.append(args.topic)
    if getattr(args, "topics", None):
        topics.extend(args.topics or [])
    # Optional: extend search terms from config
    if getattr(args, "config", None):
        try:
            cfg = load_column_config(args.config)
            extra_terms = cfg.get("search_terms", [])
            if isinstance(extra_terms, list):
                topics.extend([t for t in extra_terms if isinstance(t, str)])
        except Exception:
            pass
    topics = [t.strip() for t in topics if t and t.strip()]
    if not topics and mode == "excerpts":
        die("Provide --topic and/or --topics when using --mode excerpts")
    if not topics:
        # In full mode, allow dossiers without topic filtering.
        topics = ["selected-conversations"]

    home = home_dir(args.home)
    _, extracted_dir, dossiers_dir = ensure_layout(home)
    root = (
        Path(args.root).expanduser().resolve()
        if args.root
        else default_root(extracted_dir)
    )

    data_file = find_conversations_json(root)
    if not data_file:
        die(f"No JSON found under {root}")
    data = load_json(data_file)
    convs = normalize_conversations(data)

    wanted: List[str] = []
    if args.ids:
        wanted.extend(args.ids)
    if args.ids_file:
        p = Path(args.ids_file).expanduser().resolve()
        if not p.exists():
            die(f"IDs file not found: {p}")
        wanted.extend(
            [
                ln.strip()
                for ln in p.read_text(encoding="utf-8").splitlines()
                if ln.strip()
            ]
        )
    wanted = [w.strip() for w in wanted if w.strip()]
    if not wanted:
        die("Provide --ids and/or --ids-file")

    # Determine output formats (default to txt)
    formats = [f.lower() for f in (getattr(args, "format", None) or [])]

    # Load patterns from file if provided
    patterns = None
    if getattr(args, "patterns_file", None):
        try:
            pf = Path(args.patterns_file).expanduser().resolve()
            if pf.exists():
                patterns = [
                    ln.strip()
                    for ln in pf.read_text(encoding="utf-8").splitlines()
                    if ln.strip()
                ]
        except Exception as e:
            print(f"WARNING: Failed to load patterns file: {e}", file=sys.stderr)

    split = bool(getattr(args, "split", False))
    dedup = bool(getattr(args, "dedup", True))
    used_links_file = getattr(args, "used_links_file", None)
    config_file = getattr(args, "config", None)
    name = getattr(args, "name", None)

    out_path = build_combined_dossier(
        topics=topics,
        mode=mode,
        context=context,
        root=root,
        dossiers_dir=dossiers_dir,
        wanted_ids=wanted,
        convs=convs,
        formats=formats,
        split=split,
        patterns=patterns,
        dedup=dedup,
        used_links_file=used_links_file,
        config_file=config_file,
        name=name,
    )
    print(out_path)


def cmd_recent(args: argparse.Namespace) -> None:
    """
    Show the N most recent conversations and let you select interactively.
    Like 'quick' but without keyword filtering — just by recency.
    """
    count = int(args.count)
    if count < 1:
        die("Count must be at least 1.")

    home = home_dir(args.home)
    zips_dir, extracted_dir, dossiers_dir = ensure_layout(home)

    # Ensure we have an extracted/latest
    has_any_extracted = any(
        p.is_dir() and p.name != "latest" for p in extracted_dir.iterdir()
    )
    if not has_any_extracted:
        # auto-extract newest zip
        zpath = newest_zip(zips_dir)
        out_dir = extracted_dir / zpath.stem
        out_dir.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(zpath, "r") as zf:
            zf.extractall(out_dir)
        refresh_latest_symlink(extracted_dir, out_dir)
    else:
        newest = newest_extracted(extracted_dir)
        refresh_latest_symlink(extracted_dir, newest)

    root = (
        Path(args.root).expanduser().resolve()
        if args.root
        else default_root(extracted_dir)
    )

    data_file = find_conversations_json(root)
    if not data_file:
        die(f"No JSON found under {root}")
    data = load_json(data_file)
    convs = normalize_conversations(data)

    # Sort by create_time descending (newest first), then take top N
    convs_with_time: List[Tuple[Any, float]] = []
    for c in convs:
        cid, title = conv_id_and_title(c)
        if cid:
            ctime = float(c.get("create_time") or 0.0)
            convs_with_time.append((c, ctime))

    convs_with_time.sort(key=lambda x: x[1], reverse=True)
    recent_convs = convs_with_time[:count]

    if not recent_convs:
        die("No conversations found.")

    # Build matches list (id, title, create_time) for selection
    matches: List[Tuple[str, str, float]] = []
    for c, ctime in recent_convs:
        cid, title = conv_id_and_title(c)
        matches.append((cid, title or "", ctime))

    # For display, reverse so oldest of the N is #1 and newest is #N (chronological within the window)
    # Actually, let's keep newest at top (#1) for intuitive "most recent first"
    # matches is already newest-first from the sort above

    slug = f"recent_{count}"
    all_ids_path = dossiers_dir / f"ids__{slug}.tsv"
    selected_ids_path = dossiers_dir / f"selected_ids__{slug}.txt"

    # Write all matches
    all_lines = [f"{cid}\t{title}\n" for (cid, title, _) in matches]
    all_ids_path.write_text("".join(all_lines), encoding="utf-8")

    # Print numbered list
    print(f"\n=== {count} Most Recent Conversations ===\n")
    for i, (cid, title, ctime) in enumerate(matches, start=1):
        print(f"{i:>3}. {cid}\t{title}\t{ts_to_local_str(ctime)}")

    print(f"\nSaved full list to: {all_ids_path}")

    # Selection logic (simplified from cmd_quick)
    if args.all:
        picked = list(range(1, len(matches) + 1))
    else:

        def parse_selection_text(raw_text: str) -> Tuple[List[int], List[str]]:
            tokens = [t for t in re.split(r"[,\s]+", raw_text) if t]
            picked_local: List[int] = []
            warnings: List[str] = []
            id_to_index = {cid: idx for idx, (cid, _, _) in enumerate(matches, start=1)}

            for tok in tokens:
                if re.match(r"^\d+-\d+$", tok):
                    a, b = tok.split("-", 1)
                    a_i, b_i = int(a), int(b)
                    if a_i > b_i:
                        a_i, b_i = b_i, a_i
                    a_i = max(1, a_i)
                    b_i = min(len(matches), b_i)
                    if a_i > len(matches) or b_i < 1:
                        warnings.append(f"Range out of bounds: {tok}")
                        continue
                    for n in range(a_i, b_i + 1):
                        picked_local.append(n)
                    continue

                if tok.isdigit():
                    n = int(tok)
                    if 1 <= n <= len(matches):
                        picked_local.append(n)
                    else:
                        warnings.append(f"Selection number out of range: {n}")
                    continue

                # treat as raw ID
                if tok in id_to_index:
                    picked_local.append(id_to_index[tok])
                else:
                    warnings.append(f"Unknown ID in selection: {tok}")

            return picked_local, warnings

        stdin_is_tty = sys.stdin.isatty()
        if not stdin_is_tty:
            try:
                raw = sys.stdin.read()
            except Exception:
                die("Failed to read stdin for selection.")
            if not raw or not raw.strip():
                die("No selection provided on stdin.")
            picked, warnings = parse_selection_text(raw)
            if warnings:
                for w in warnings:
                    print(f"WARNING: {w}", file=sys.stderr)
        else:
            while True:
                try:
                    raw = input(
                        "\nPick by number (e.g. 1 3 7), range (1-5), or 'all': "
                    ).strip()
                except (KeyboardInterrupt, EOFError):
                    print("\nSelection cancelled.")
                    return
                if not raw:
                    die("No selection provided.")
                if raw.lower() == "all":
                    picked = list(range(1, len(matches) + 1))
                    break

                picked, warnings = parse_selection_text(raw)
                if not warnings:
                    break
                for w in warnings:
                    print(f"WARNING: {w}", file=sys.stderr)
                try:
                    correction = input(
                        "\nErrors detected. Enter corrected selection, or ENTER to accept valid picks: "
                    )
                except (KeyboardInterrupt, EOFError):
                    print("\nSelection cancelled.")
                    return
                if correction is None:
                    print("\nSelection cancelled.")
                    return
                correction = correction.strip()
                if correction == "":
                    if not picked:
                        print("No valid selections; please try again.", file=sys.stderr)
                        continue
                    break
                raw = correction

    picked = sorted(set(picked))
    if not picked:
        die("No valid selections were parsed.")
    wanted_ids = [matches[i - 1][0] for i in picked]

    selected_ids_path.write_text("\n".join(wanted_ids) + "\n", encoding="utf-8")
    print(f"\nSaved selected IDs to: {selected_ids_path}")

    # Build dossier
    formats = [f.lower() for f in (getattr(args, "format", None) or [])]
    patterns = None
    if getattr(args, "patterns_file", None):
        try:
            pf = Path(args.patterns_file).expanduser().resolve()
            if pf.exists():
                patterns = [
                    ln.strip()
                    for ln in pf.read_text(encoding="utf-8").splitlines()
                    if ln.strip()
                ]
        except Exception as e:
            print(f"WARNING: Failed to load patterns file: {e}", file=sys.stderr)

    split = bool(getattr(args, "split", False))
    dedup = bool(getattr(args, "dedup", True))
    used_links_file = getattr(args, "used_links_file", None)
    config_file = getattr(args, "config", None)
    mode = getattr(args, "mode", None) or "full"
    context = int(getattr(args, "context", 2))
    name = getattr(args, "name", None)

    out_path = build_combined_dossier(
        topics=[f"recent-{count}"],
        mode=mode,
        context=context,
        root=root,
        dossiers_dir=dossiers_dir,
        wanted_ids=wanted_ids,
        convs=convs,
        formats=formats,
        split=split,
        patterns=patterns,
        dedup=dedup,
        used_links_file=used_links_file,
        config_file=config_file,
        name=name,
    )
    print(f"\nWrote dossier: {out_path}")


def cmd_quick(args: argparse.Namespace) -> None:
    """
    One command that:
      - extracts newest zip (if nothing extracted yet)
      - searches titles by topic(s)
      - lets you select by number (or 'all')
      - builds dossier from the selected IDs
    """
    topics = [t.strip() for t in args.topics if t.strip()]
    if getattr(args, "config", None):
        try:
            cfg = load_column_config(args.config)
            extra_terms = cfg.get("search_terms", [])
            if isinstance(extra_terms, list):
                topics.extend([t for t in extra_terms if isinstance(t, str)])
        except Exception:
            pass
    if not topics:
        die("Provide at least one topic.")

    mode = args.mode
    context = int(args.context)
    recent_count = getattr(args, "recent_count", None)
    days_count = getattr(args, "days_count", None)
    if recent_count is not None and recent_count < 1:
        die("--recent must be >= 1")
    if days_count is not None and days_count < 1:
        die("--days must be >= 1")

    home = home_dir(args.home)
    zips_dir, extracted_dir, dossiers_dir = ensure_layout(home)

    # Ensure we have an extracted/latest
    has_any_extracted = any(
        p.is_dir() and p.name != "latest" for p in extracted_dir.iterdir()
    )
    if not has_any_extracted:
        # auto-extract newest zip
        zpath = newest_zip(zips_dir)
        out_dir = extracted_dir / zpath.stem
        out_dir.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(zpath, "r") as zf:
            zf.extractall(out_dir)
        refresh_latest_symlink(extracted_dir, out_dir)
    else:
        # refresh latest to newest extracted dir (keeps things consistent)
        newest = newest_extracted(extracted_dir)
        refresh_latest_symlink(extracted_dir, newest)

    root = (
        Path(args.root).expanduser().resolve()
        if args.root
        else default_root(extracted_dir)
    )

    data_file = find_conversations_json(root)
    if not data_file:
        die(f"No JSON found under {root}")
    data = load_json(data_file)
    convs = normalize_conversations(data)
    if days_count is not None:
        now_ts = datetime.now(tz=timezone.utc).timestamp()
        cutoff_ts = now_ts - (days_count * 86400.0)
        convs = [
            c
            for c in convs
            if c.get("id") and float(c.get("create_time") or 0.0) >= cutoff_ts
        ]
    if recent_count is not None:
        convs_with_time: List[Tuple[Any, float]] = []
        for c in convs:
            cid, _ = conv_id_and_title(c)
            if not cid:
                continue
            ctime = float(c.get("create_time") or 0.0)
            convs_with_time.append((c, ctime))
        convs_with_time.sort(key=lambda x: x[1], reverse=True)
        convs = [c for c, _ in convs_with_time[:recent_count]]

    needles = [t.lower() for t in topics]
    and_terms = bool(args.and_terms)
    where = getattr(args, "where", "title")
    topic_re = compile_topic_pattern(topics)

    matches: List[Tuple[str, str, float]] = []  # (id, title, create_time)
    for c in convs:
        cid, title = conv_id_and_title(c)
        if not cid:
            continue
        matched = False
        if where in ("title", "all"):
            t = (title or "").lower()
            ok_title = (
                all(n in t for n in needles)
                if and_terms
                else any(n in t for n in needles)
            )
            if ok_title:
                matched = True
        if not matched and where in ("messages", "all"):
            if conversation_matches_text(c, topic_re):
                matched = True
        if matched:
            ctime = float(c.get("create_time") or 0.0)
            matches.append((cid, title or "", ctime))

    if not matches:
        die("No conversation titles matched those topic terms.")

    # Sort by conversation create_time (sane for selection)
    matches.sort(key=lambda x: x[2])

    slug = safe_slug("_".join(topics))
    all_ids_path = dossiers_dir / f"ids__{slug}.tsv"
    selected_ids_path = dossiers_dir / f"selected_ids__{slug}.txt"

    # Write all matches (for manual copy if you want)
    all_lines = [f"{cid}\t{title}\n" for (cid, title, _) in matches]
    all_ids_path.write_text("".join(all_lines), encoding="utf-8")

    # Print numbered list
    for i, (cid, title, ctime) in enumerate(matches, start=1):
        colored_title = _colorize_title_with_topics(title or "", topics)
        print(f"{i:>3}. {cid}\t{colored_title}\t{ts_to_local_str(ctime)}")

    print(f"\nSaved full match list to: {all_ids_path}")

    def parse_selection_text(raw_text: str) -> Tuple[List[int], List[str]]:
        tokens = [t for t in re.split(r"[,\s]+", raw_text) if t]
        picked_local: List[int] = []
        warnings: List[str] = []
        id_to_index = {cid: idx for idx, (cid, _, _) in enumerate(matches, start=1)}

        for tok in tokens:
            if tok.startswith("@"):
                # load IDs from file
                path = Path(tok[1:]).expanduser()
                if not path.exists():
                    warnings.append(f"IDs file not found: {path}")
                    continue
                for ln in path.read_text(encoding="utf-8").splitlines():
                    ln = ln.strip()
                    if not ln:
                        continue
                    if ln in id_to_index:
                        picked_local.append(id_to_index[ln])
                    else:
                        # allow raw numeric indices in file
                        if ln.isdigit():
                            n = int(ln)
                            if 1 <= n <= len(matches):
                                picked_local.append(n)
                            else:
                                warnings.append(
                                    f"Selection number out of range in file: {n}"
                                )
                        else:
                            warnings.append(f"Unknown ID in file: {ln}")
                continue

            if re.match(r"^\d+-\d+$", tok):
                a, b = tok.split("-", 1)
                a_i = int(a)
                b_i = int(b)
                if a_i > b_i:
                    a_i, b_i = b_i, a_i
                # clamp to valid range
                a_i = max(1, a_i)
                b_i = min(len(matches), b_i)
                if a_i > len(matches) or b_i < 1:
                    warnings.append(f"Range out of bounds: {tok}")
                    continue
                for n in range(a_i, b_i + 1):
                    picked_local.append(n)
                continue

            if tok.isdigit():
                n = int(tok)
                if 1 <= n <= len(matches):
                    picked_local.append(n)
                else:
                    warnings.append(f"Selection number out of range: {n}")
                continue

            # treat as raw ID
            if tok in id_to_index:
                picked_local.append(id_to_index[tok])
            else:
                warnings.append(f"Unknown ID in selection: {tok}")

        return picked_local, warnings

    # Non-interactive selection via --ids-file: read file and parse selections.
    if getattr(args, "ids_file", None):
        p = Path(args.ids_file).expanduser().resolve()
        if not p.exists():
            die(f"IDs file not found: {p}")
        raw = "\n".join(
            [
                ln.strip()
                for ln in p.read_text(encoding="utf-8").splitlines()
                if ln.strip()
            ]
        )
        picked, warnings = parse_selection_text(raw)
        if warnings:
            for w in warnings:
                print(f"WARNING: {w}", file=sys.stderr)
    elif args.all:
        picked = list(range(1, len(matches) + 1))
    else:
        stdin_is_tty = sys.stdin.isatty()
        if not stdin_is_tty:
            # Non-interactive stdin: read once and proceed (warnings printed)
            try:
                raw = sys.stdin.read()
            except Exception:
                die("Failed to read stdin for selection.")
            if not raw or not raw.strip():
                die("No selection provided on stdin.")
            picked, warnings = parse_selection_text(raw)
            if warnings:
                for w in warnings:
                    print(f"WARNING: {w}", file=sys.stderr)
        else:
            # Interactive: allow correction loop when there are warnings
            while True:
                try:
                    raw = input(
                        "\nPick by number (e.g. 1 3 7), or 'all', or paste IDs: "
                    ).strip()
                except (KeyboardInterrupt, EOFError):
                    print("\nSelection cancelled.")
                    return
                if not raw:
                    die("No selection provided.")
                if raw.lower() == "all":
                    picked = list(range(1, len(matches) + 1))
                    warnings = []
                    break

                picked, warnings = parse_selection_text(raw)

                if not warnings:
                    break  # good selection, proceed

                # There are warnings: show them and allow user to correct or accept partial selection
                for w in warnings:
                    print(f"WARNING: {w}", file=sys.stderr)
                try:
                    correction = input(
                        "\nErrors detected. Enter corrected selection, or press ENTER to accept valid selections and continue: "
                    )
                except (KeyboardInterrupt, EOFError):
                    print("\nSelection cancelled.")
                    return
                if correction is None:
                    print("\nSelection cancelled.")
                    return
                correction = correction.strip()
                if correction == "":
                    # accept current valid picks (if any)
                    if not picked:
                        print(
                            "No valid selections to accept; please enter a corrected selection.",
                            file=sys.stderr,
                        )
                        continue
                    break
                # otherwise loop will parse the new correction text

                raw = correction
                picked, warnings = parse_selection_text(raw)
                if not warnings:
                    break
                # else loop continues and user will be prompted again

    picked = sorted(set(picked))
    if not picked:
        die("No valid selections were parsed. Check your selection numbers/IDs.")
    wanted_ids = [matches[i - 1][0] for i in picked]

    selected_ids_path.write_text("\n".join(wanted_ids) + "\n", encoding="utf-8")
    print(f"\nSaved selected IDs to: {selected_ids_path}")

    formats = [f.lower() for f in (getattr(args, "format", None) or [])]

    # Load patterns from file if provided
    patterns = None
    if getattr(args, "patterns_file", None):
        try:
            pf = Path(args.patterns_file).expanduser().resolve()
            if pf.exists():
                patterns = [
                    ln.strip()
                    for ln in pf.read_text(encoding="utf-8").splitlines()
                    if ln.strip()
                ]
        except Exception as e:
            print(f"WARNING: Failed to load patterns file: {e}", file=sys.stderr)

    split = bool(getattr(args, "split", False))
    dedup = bool(getattr(args, "dedup", True))
    used_links_file = getattr(args, "used_links_file", None)
    config_file = getattr(args, "config", None)
    name = getattr(args, "name", None)

    out_path = build_combined_dossier(
        topics=topics,
        mode=mode,
        context=context,
        root=root,
        dossiers_dir=dossiers_dir,
        wanted_ids=wanted_ids,
        convs=convs,
        formats=formats,
        split=split,
        patterns=patterns,
        dedup=dedup,
        used_links_file=used_links_file,
        config_file=config_file,
        name=name,
    )
    print(f"\nWrote dossier: {out_path}")


def cmd_index(args: argparse.Namespace) -> None:
    """(Re)build the search index from an extracted export."""
    home = home_dir(getattr(args, "home", None))
    _, extracted_dir, _ = ensure_layout(home)
    root = (
        Path(args.root).expanduser().resolve()
        if getattr(args, "root", None)
        else default_root(extracted_dir)
    )
    db_path = (
        Path(args.db).expanduser().resolve()
        if getattr(args, "db", None)
        else extracted_dir / "cgpt_index.db"
    )
    try:
        index_export(
            root,
            db_path,
            reindex=bool(getattr(args, "reindex", False)),
            show_progress=not getattr(args, "quiet", False),
        )
        print(f"Index built at: {db_path}")
    except Exception as e:
        die(f"Indexing failed: {e}")


def _add_split_flags(parser: argparse.ArgumentParser, split_help: str) -> None:
    """Add --split/--no-split flags with tri-state default for env fallback."""
    grp = parser.add_mutually_exclusive_group()
    grp.add_argument(
        "--split",
        dest="split",
        action="store_true",
        default=None,
        help=split_help,
    )
    grp.add_argument(
        "--no-split",
        dest="split",
        action="store_false",
        help="Disable split output (overrides CGPT_DEFAULT_SPLIT).",
    )


def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="cgpt.py",
        description="ChatGPT export helper (zips → extracted → dossiers).",
    )
    p.add_argument(
        "--version", action="version", version=f"%(prog)s {__version__}"
    )
    # CLI-level color control: --color / --no-color
    color_grp = p.add_mutually_exclusive_group()
    color_grp.add_argument(
        "--color", dest="color", action="store_true", help="Force-enable ANSI colors"
    )
    color_grp.add_argument(
        "--no-color",
        dest="no_color",
        action="store_true",
        help="Force-disable ANSI colors",
    )
    p.add_argument(
        "--home",
        help="Home folder containing zips/, extracted/, dossiers/. Default: $CGPT_HOME, auto-detected, or CWD",
    )
    p.add_argument(
        "--quiet",
        dest="quiet",
        action="store_true",
        help="Suppress non-error output (useful in scripts)",
    )
    p.add_argument(
        "--default-mode",
        dest="default_mode",
        choices=["full", "excerpts"],
        default=None,
        help="Set preferred default mode for dossier creation (overrides CGPT_DEFAULT_MODE)",
    )
    # If no subcommand is provided, we'll default to extracting the newest ZIP.
    sub = p.add_subparsers(dest="cmd", required=False)

    a = sub.add_parser(
        "init", help="Create/verify required folders: zips/, extracted/, dossiers/"
    )
    a.set_defaults(func=cmd_init)

    a = sub.add_parser("latest-zip", help="Print newest ZIP in zips/")
    a.set_defaults(func=cmd_latest_zip)

    a = sub.add_parser(
        "extract",
        help="Extract a ZIP into extracted/<zip_stem>/ (defaults to newest ZIP)",
    )
    a.add_argument("zip", nargs="?", help="Path to ZIP (optional)")
    a.add_argument(
        "--no-index",
        dest="no_index",
        action="store_true",
        help="Do not update the search index after extracting",
    )
    a.add_argument(
        "--reindex",
        dest="reindex",
        action="store_true",
        help="Force rebuild of the search index after extracting",
    )
    a.set_defaults(func=cmd_extract)

    a = sub.add_parser(
        "index",
        help="(Re)build the search index from an extracted export (uses FTS5 when available)",
    )
    a.add_argument(
        "--root", help="Extracted folder to scan (defaults to extracted/latest)"
    )
    a.add_argument(
        "--reindex",
        dest="reindex",
        action="store_true",
        help="Force rebuild of the search index",
    )
    a.add_argument(
        "--db",
        dest="db",
        help="Path to index DB file (defaults to extracted/cgpt_index.db)",
    )
    a.set_defaults(func=cmd_index)

    # Short alias
    a = sub.add_parser("x", help="Alias for extract")
    a.add_argument("zip", nargs="?", help="Path to ZIP (optional)")
    a.add_argument(
        "--no-index",
        dest="no_index",
        action="store_true",
        help="Do not update the search index after extracting",
    )
    a.add_argument(
        "--reindex",
        dest="reindex",
        action="store_true",
        help="Force rebuild of the search index after extracting",
    )
    a.set_defaults(func=cmd_extract)

    a = sub.add_parser("ids", help="Print id<TAB>title for all conversations")
    a.add_argument(
        "--root", help="Extracted folder to scan (defaults to extracted/latest)"
    )
    a.set_defaults(func=cmd_ids)

    # Short alias
    a = sub.add_parser("i", help="Alias for ids")
    a.add_argument(
        "--root", help="Extracted folder to scan (defaults to extracted/latest)"
    )
    a.set_defaults(func=cmd_ids)

    a = sub.add_parser(
        "find", help="Find conversations whose titles match query (case-insensitive)"
    )
    a.add_argument("query")
    a.add_argument(
        "--root", help="Extracted folder to scan (defaults to extracted/latest)"
    )
    a.set_defaults(func=cmd_find)

    # Short alias
    a = sub.add_parser("f", help="Alias for find")
    a.add_argument("query")
    a.add_argument(
        "--root", help="Extracted folder to scan (defaults to extracted/latest)"
    )
    a.set_defaults(func=cmd_find)

    a = sub.add_parser(
        "search", help="Search in titles and/or message text (case-insensitive)"
    )
    a.add_argument("query", nargs="?", default=None)
    a.add_argument(
        "where",
        nargs="?",
        choices=["title", "messages", "all"],
        help="Optional positional: where to search (title|messages|all)",
    )
    a.add_argument(
        "--terms",
        nargs="+",
        help="One or more search terms (use with --and to require all)",
    )
    a.add_argument(
        "--and",
        dest="and_terms",
        action="store_true",
        help="Require ALL terms to match (default is OR)",
    )
    a.add_argument(
        "--where",
        dest="where_opt",
        choices=["title", "messages", "all"],
        default=None,
        help="Where to search: title (default), messages, or all",
    )
    a.add_argument(
        "--root", help="Extracted folder to scan (defaults to extracted/latest)"
    )
    a.set_defaults(func=cmd_search)

    a = sub.add_parser(
        "make-dossiers", help="Write one or more formats per selected conversation ID"
    )
    a.add_argument(
        "--root", help="Extracted folder to scan (defaults to extracted/latest)"
    )
    a.add_argument("--ids-file", help="Text file with one id per line")
    a.add_argument("--ids", nargs="*", help="One or more IDs")
    a.add_argument(
        "--format",
        nargs="+",
        choices=["txt", "md", "docx"],
        default=["txt"],
        help=(
            "Output format(s) for per-conversation dossiers (default: txt). "
            "Examples: --format txt  # plain text (default); "
            "--format md docx  # produce Markdown and Word .docx; "
            "--format txt md docx  # produce all three"
        ),
    )
    a.set_defaults(func=cmd_make_dossiers)

    a = sub.add_parser(
        "build-dossier",
        help="Build a single combined dossier with time + branch nesting",
    )
    a.add_argument("--topic", help="Single topic keyword (for excerpts mode)")
    a.add_argument(
        "--topics", nargs="*", help="One or more topics (OR logic) (for excerpts mode)"
    )
    a.add_argument(
        "--mode",
        choices=["full", "excerpts"],
        default=None,
        help="full = include everything; excerpts = topic-only + context",
    )
    a.add_argument(
        "--context",
        default=2,
        help="In excerpts mode, include +/- N messages around matches",
    )
    a.add_argument(
        "--root", help="Extracted folder to scan (defaults to extracted/latest)"
    )
    a.add_argument("--ids-file", help="Text file with one id per line")
    a.add_argument("--ids", nargs="*", help="One or more IDs")
    a.add_argument(
        "--format",
        nargs="+",
        choices=["txt", "md", "docx"],
        help=(
            "One or more output formats: txt (default), md, docx. "
            "Examples: --format txt; --format md docx; --format txt md docx"
        ),
        default=["txt"],
    )
    _add_split_flags(
        a,
        "Generate two TXT files: dossier_raw.txt (full) and dossier_raw__working.txt (cleaned, deduplicated, deliverables-only).",
    )
    a.add_argument(
        "--dedup",
        action="store_true",
        default=True,
        help="Enable deduplication in working output (default: True)",
    )
    a.add_argument(
        "--no-dedup",
        dest="dedup",
        action="store_false",
        help="Disable deduplication in working output",
    )
    a.add_argument(
        "--patterns-file",
        help="Path to file with deliverable patterns (one per line). Default patterns: ##, Constraint, Draft, Decision, Output, Result",
    )
    a.add_argument(
        "--used-links-file",
        help="Path to file with URLs already used in drafts (one per line). These will be prioritized in source lists.",
    )
    a.add_argument(
        "--config",
        help="Path to column config file (JSON) for segment filtering and control layer generation",
    )
    a.add_argument(
        "--name",
        help="Project name for organizing output. Creates dossiers/{name}/ subfolder.",
    )
    a.set_defaults(func=cmd_build_dossier)

    # Short alias
    a = sub.add_parser("d", help="Alias for build-dossier")
    a.add_argument("--topic", help="Single topic keyword (for excerpts mode)")
    a.add_argument(
        "--topics", nargs="*", help="One or more topics (OR logic) (for excerpts mode)"
    )
    a.add_argument("--mode", choices=["full", "excerpts"], default=None)
    a.add_argument("--context", default=2)
    a.add_argument(
        "--root", help="Extracted folder to scan (defaults to extracted/latest)"
    )
    a.add_argument("--ids-file", help="Text file with one id per line")
    a.add_argument("--ids", nargs="*", help="One or more IDs")
    a.add_argument(
        "--format",
        nargs="+",
        choices=["txt", "md", "docx"],
        default=["txt"],
    )
    _add_split_flags(
        a,
        "Generate two TXT files: dossier_raw.txt (full) and dossier_raw__working.txt (cleaned, deduplicated, deliverables-only).",
    )
    a.add_argument(
        "--dedup",
        action="store_true",
        default=True,
        help="Enable deduplication in working output (default: True)",
    )
    a.add_argument(
        "--no-dedup",
        dest="dedup",
        action="store_false",
        help="Disable deduplication in working output",
    )
    a.add_argument(
        "--patterns-file",
        help="Path to file with deliverable patterns (one per line). Default patterns: ##, Constraint, Draft, Decision, Output, Result",
    )
    a.add_argument(
        "--used-links-file",
        help="Path to file with URLs already used in drafts (one per line). These will be prioritized in source lists.",
    )
    a.add_argument(
        "--config",
        help="Path to column config file (JSON) for segment filtering and control layer generation",
    )
    a.add_argument(
        "--name",
        help="Project name for organizing output. Creates dossiers/{name}/ subfolder.",
    )
    a.set_defaults(func=cmd_build_dossier)

    # One-shot command for minimal typing
    a = sub.add_parser(
        "quick", help="Extract (if needed) → find by title → pick IDs → build dossier"
    )
    a.add_argument(
        "topics",
        nargs="+",
        help="One or more topic terms to match in conversation titles",
    )
    a.add_argument(
        "--and",
        dest="and_terms",
        action="store_true",
        help="Require ALL terms to appear in title (default is OR)",
    )
    a.add_argument("--mode", choices=["full", "excerpts"], default=None)
    a.add_argument("--context", default=2)
    a.add_argument("--all", action="store_true", help="Select all matches (no prompt)")
    a.add_argument(
        "--where",
        choices=["title", "messages", "all"],
        default="title",
        help="Where to search when matching topics: title (default), messages, or all",
    )
    quick_recency = a.add_mutually_exclusive_group()
    quick_recency.add_argument(
        "--recent",
        dest="recent_count",
        type=int,
        default=None,
        help="Limit quick matching to the N most recent conversations before applying topic filters",
    )
    quick_recency.add_argument(
        "--days",
        dest="days_count",
        type=int,
        default=None,
        help="Limit quick matching to conversations created in the last N days before applying topic filters",
    )
    a.add_argument(
        "--root", help="Extracted folder to scan (defaults to extracted/latest)"
    )
    a.add_argument("--ids-file", help="Text file with one id per line")
    a.add_argument(
        "--format",
        nargs="+",
        choices=["txt", "md", "docx"],
        default=["txt"],
        help=(
            "Output formats for dossier (default: txt). "
            "Examples: python3 cgpt.py quick --format txt 'term1' 'term2'; "
            "python3 cgpt.py quick --format md docx 'research' 'analysis'"
        ),
    )
    _add_split_flags(
        a,
        "Generate two TXT files: dossier_raw.txt (full) and dossier_raw__working.txt (cleaned, deduplicated, deliverables-only).",
    )
    a.add_argument(
        "--dedup",
        action="store_true",
        default=True,
        help="Enable deduplication in working output (default: True)",
    )
    a.add_argument(
        "--no-dedup",
        dest="dedup",
        action="store_false",
        help="Disable deduplication in working output",
    )
    a.add_argument(
        "--patterns-file",
        help="Path to file with deliverable patterns (one per line). Default patterns: ##, Constraint, Draft, Decision, Output, Result",
    )
    a.add_argument(
        "--used-links-file",
        help="Path to file with URLs already used in drafts (one per line).",
    )
    a.add_argument(
        "--config",
        help="Path to column config file (JSON) for segment filtering and control layer generation",
    )
    a.add_argument(
        "--name",
        help="Project name for organizing output. Creates dossiers/{name}/ subfolder.",
    )
    a.set_defaults(func=cmd_quick)

    # Very short alias
    a = sub.add_parser("q", help="Alias for quick")
    a.add_argument("topics", nargs="+")
    a.add_argument("--and", dest="and_terms", action="store_true")
    a.add_argument("--mode", choices=["full", "excerpts"], default=None)
    a.add_argument("--context", default=2)
    a.add_argument("--all", action="store_true")
    a.add_argument(
        "--where",
        choices=["title", "messages", "all"],
        default="title",
        help="Where to search when matching topics: title (default), messages, or all",
    )
    quick_recency = a.add_mutually_exclusive_group()
    quick_recency.add_argument(
        "--recent",
        dest="recent_count",
        type=int,
        default=None,
        help="Limit quick matching to the N most recent conversations before applying topic filters",
    )
    quick_recency.add_argument(
        "--days",
        dest="days_count",
        type=int,
        default=None,
        help="Limit quick matching to conversations created in the last N days before applying topic filters",
    )
    a.add_argument(
        "--root", help="Extracted folder to scan (defaults to extracted/latest)"
    )
    a.add_argument("--ids-file", help="Text file with one id per line")
    a.add_argument(
        "--format",
        nargs="+",
        choices=["txt", "md", "docx"],
        default=["txt"],
    )
    _add_split_flags(
        a,
        "Generate two TXT files: dossier_raw.txt (full) and dossier_raw__working.txt (cleaned, deduplicated, deliverables-only).",
    )
    a.add_argument(
        "--dedup",
        action="store_true",
        default=True,
        help="Enable deduplication in working output (default: True)",
    )
    a.add_argument(
        "--no-dedup",
        dest="dedup",
        action="store_false",
        help="Disable deduplication in working output",
    )
    a.add_argument(
        "--patterns-file",
        help="Path to file with deliverable patterns (one per line). Default patterns: ##, Constraint, Draft, Decision, Output, Result",
    )
    a.add_argument(
        "--used-links-file",
        help="Path to file with URLs already used in drafts (one per line).",
    )
    a.add_argument(
        "--config",
        help="Path to column config file (JSON) for segment filtering and control layer generation",
    )
    a.add_argument(
        "--name",
        help="Project name for organizing output. Creates dossiers/{name}/ subfolder.",
    )
    a.set_defaults(func=cmd_quick)

    # Recent command: browse N most recent conversations interactively
    a = sub.add_parser(
        "recent",
        help="Show the N most recent conversations and select interactively",
    )
    a.add_argument(
        "count",
        type=int,
        nargs="?",
        default=30,
        help="Number of recent conversations to show (default: 30)",
    )
    a.add_argument("--all", action="store_true", help="Select all shown (no prompt)")
    a.add_argument(
        "--root", help="Extracted folder to scan (defaults to extracted/latest)"
    )
    a.add_argument(
        "--format",
        nargs="+",
        choices=["txt", "md", "docx"],
        default=["txt"],
        help="Output format(s) for dossier (default: txt)",
    )
    _add_split_flags(a, "Generate both raw and working TXT files.")
    a.add_argument(
        "--dedup",
        action="store_true",
        default=True,
        help="Enable deduplication in working output (default: True)",
    )
    a.add_argument(
        "--no-dedup",
        dest="dedup",
        action="store_false",
        help="Disable deduplication in working output",
    )
    a.add_argument(
        "--patterns-file",
        help="Path to file with deliverable patterns (one per line)",
    )
    a.add_argument(
        "--used-links-file",
        help="Path to file with URLs already used in drafts (one per line)",
    )
    a.add_argument(
        "--config",
        help="Path to column config file (JSON)",
    )
    a.add_argument(
        "--name",
        help="Project name for organizing output. Creates dossiers/{name}/ subfolder.",
    )
    a.add_argument("--mode", choices=["full", "excerpts"], default=None)
    a.add_argument("--context", default=2)
    a.set_defaults(func=cmd_recent)

    # Short alias for recent
    a = sub.add_parser("r", help="Alias for recent")
    a.add_argument("count", type=int, nargs="?", default=30)
    a.add_argument("--all", action="store_true")
    a.add_argument("--root")
    a.add_argument(
        "--format", nargs="+", choices=["txt", "md", "docx"], default=["txt"]
    )
    _add_split_flags(a, "Generate both raw and working TXT files.")
    a.add_argument("--dedup", action="store_true", default=True)
    a.add_argument("--no-dedup", dest="dedup", action="store_false")
    a.add_argument("--patterns-file")
    a.add_argument("--used-links-file")
    a.add_argument("--config")
    a.add_argument("--name", help="Project name for organizing output.")
    a.add_argument("--mode", choices=["full", "excerpts"], default=None)
    a.add_argument("--context", default=2)
    a.set_defaults(func=cmd_recent)

    return p


def main() -> None:
    args = build_parser().parse_args()

    # Honor CLI color flags (override env and auto-detect). Must set before any coloring.
    global _CLI_COLOR_OVERRIDE
    if getattr(args, "color", False):
        _CLI_COLOR_OVERRIDE = True
    elif getattr(args, "no_color", False):
        _CLI_COLOR_OVERRIDE = False

    # Default behavior: if no subcommand provided, extract newest ZIP in `zips/`.
    if not getattr(args, "cmd", None):
        # Ensure args has a `zip` attribute for cmd_extract (it expects args.zip)
        if not hasattr(args, "zip"):
            setattr(args, "zip", None)
        # Respect global quiet flag when default-extract
        if getattr(args, "quiet", False):
            setattr(args, "quiet", True)
        cmd_extract(args)
        return
    # Resolve global/default mode preference: CLI > env CGPT_DEFAULT_MODE > builtin 'full'
    effective_default_mode = None
    if getattr(args, "default_mode", None):
        effective_default_mode = args.default_mode
    else:
        env_mode = os.environ.get("CGPT_DEFAULT_MODE")
        if env_mode and env_mode.lower() in ("full", "excerpts"):
            effective_default_mode = env_mode.lower()
    if effective_default_mode is None:
        effective_default_mode = "full"

    # If the chosen subcommand has a `mode` attribute that wasn't explicitly provided
    # (we set subparser defaults to None), fill it from the effective default.
    if hasattr(args, "mode") and getattr(args, "mode") is None:
        setattr(args, "mode", effective_default_mode)
    # Resolve split default from env when subcommand supports split and CLI did not set it.
    # Priority: CLI --split/--no-split > CGPT_DEFAULT_SPLIT > builtin False.
    if hasattr(args, "split") and getattr(args, "split") is None:
        env_split = _parse_env_bool("CGPT_DEFAULT_SPLIT")
        setattr(args, "split", env_split if env_split is not None else False)
    args.func(args)


if __name__ == "__main__":
    main()
